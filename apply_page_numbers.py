#!/usr/bin/env python3
"""Apply page numbering to FYP Report - Final.docx"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

INPUT  = 'Documentation/FYP Report - Final.docx'
OUTPUT = 'Documentation/FYP Report - Final-numbered.docx'

doc = Document(INPUT)
XML_NS = 'http://www.w3.org/XML/1998/namespace'

# ── helpers ────────────────────────────────────────────────────────────────

def _pPr(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    return pPr

def clear_para(p_el):
    """Remove runs/fields, keep pPr."""
    for child in list(p_el):
        if child.tag != qn('w:pPr'):
            p_el.remove(child)

def add_page_field(p_el):
    """Append centred PAGE field to a paragraph element."""
    pPr = _pPr(p_el)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    r1 = OxmlElement('w:r'); fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin'); r1.append(fc1); p_el.append(r1)

    r2 = OxmlElement('w:r'); it = OxmlElement('w:instrText')
    it.set(f'{{{XML_NS}}}space', 'preserve'); it.text = ' PAGE '
    r2.append(it); p_el.append(r2)

    r3 = OxmlElement('w:r'); fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate'); r3.append(fc2); p_el.append(r3)

    r4 = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '1'
    r4.append(t); p_el.append(r4)

    r5 = OxmlElement('w:r'); fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end'); r5.append(fc3); p_el.append(r5)

def _get_first_para(ftr_obj):
    if ftr_obj.paragraphs:
        return ftr_obj.paragraphs[0]._p
    p = OxmlElement('w:p')
    ftr_obj._element.append(p)
    return p

def setup_footer(section, show_num):
    ftr = section.footer
    ftr.is_linked_to_previous = False
    p_el = _get_first_para(ftr)
    clear_para(p_el)
    if show_num:
        add_page_field(p_el)

def setup_first_footer(section, show_num):
    ftr = section.first_page_footer
    ftr.is_linked_to_previous = False
    p_el = _get_first_para(ftr)
    clear_para(p_el)
    if show_num:
        add_page_field(p_el)

def set_pgNumType(sectPr, fmt, start=None):
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pg = OxmlElement('w:pgNumType')
    pg.set(qn('w:fmt'), fmt)
    if start is not None:
        pg.set(qn('w:start'), str(start))
    sectPr.append(pg)

def enable_titlePg(sectPr):
    if sectPr.find(qn('w:titlePg')) is None:
        sectPr.append(OxmlElement('w:titlePg'))

def insert_sect_break(para, ref_sectPr):
    """Insert a minimal sectPr into para's pPr (makes para the last of a section)."""
    p_el = para._p
    pPr = _pPr(p_el)
    old = pPr.find(qn('w:sectPr'))
    if old is not None:
        pPr.remove(old)
    new = OxmlElement('w:sectPr')
    for tag in [qn('w:pgSz'), qn('w:pgMar')]:
        el = ref_sectPr.find(tag)
        if el is not None:
            new.append(deepcopy(el))
    pPr.append(new)
    return new

# ── STEP 1: split Section 0 → title page | front matter ───────────────────

paragraphs = doc.paragraphs
ref_sectPr = doc.sections[0]._sectPr

title_end_idx = next(
    i for i, p in enumerate(paragraphs)
    if 'June' in p.text and '2026' in p.text
)
print(f"Title page ends at para [{title_end_idx}]: {paragraphs[title_end_idx].text[:40]}")

insert_sect_break(paragraphs[title_end_idx], ref_sectPr)

sections = doc.sections
print(f"Sections after split: {len(sections)}")
# Expected 17 (was 16, now 17)

# ── STEP 2: global cleanup — strip inherited pgNumType / titlePg ──────────
# The original document had its own page-number resets on many sections.
# Wipe them all first; only the sections I explicitly configure get them.

for sec in doc.sections:
    sp = sec._sectPr
    for el in list(sp.findall(qn('w:pgNumType'))):
        sp.remove(el)
    tpg = sp.find(qn('w:titlePg'))
    if tpg is not None:
        sp.remove(tpg)

sections = doc.sections   # re-read after cleanup

# ── STEP 3: configure each section ────────────────────────────────────────

s = sections   # shorthand

# S0 — title page: NO footer
setup_footer(s[0], show_num=False)

# S1 — front matter (Approval → LOA): Roman i, ii, iii …
setup_footer(s[1], show_num=True)
set_pgNumType(s[1]._sectPr, 'lowerRoman', start=1)

# S2 — Chapter 1 (title + content in one section): Arabic, titlePg
enable_titlePg(s[2]._sectPr)
setup_first_footer(s[2], show_num=False)   # Ch1 title page: hidden
setup_footer(s[2], show_num=True)
set_pgNumType(s[2]._sectPr, 'decimal', start=1)

# S3 — Chapter 2 title page only: no number
setup_footer(s[3], show_num=False)

# S4 — Chapter 2 content: Arabic (continue)
setup_footer(s[4], show_num=True)

# S5 — Chapter 3 title page: no number
setup_footer(s[5], show_num=False)

# S6, S7 — spurious empty sections
setup_footer(s[6], show_num=False)
setup_footer(s[7], show_num=False)

# S8 — Chapter 3 content
setup_footer(s[8], show_num=True)

# S9 — Chapter 4 title page: no number
setup_footer(s[9], show_num=False)

# S10 — Chapter 4 content
setup_footer(s[10], show_num=True)

# S11 — Chapter 5 title page: no number
setup_footer(s[11], show_num=False)

# S12 — Chapter 5 content
setup_footer(s[12], show_num=True)

# S13 — Chapter 6 (title + content): titlePg
enable_titlePg(s[13]._sectPr)
setup_first_footer(s[13], show_num=False)
setup_footer(s[13], show_num=True)

# S14 — Chapter 7 (title + content): titlePg
enable_titlePg(s[14]._sectPr)
setup_first_footer(s[14], show_num=False)
setup_footer(s[14], show_num=True)

# S15 — References
setup_footer(s[15], show_num=True)

# S16 — final body section (if present)
if len(s) > 16:
    setup_footer(s[16], show_num=True)

# ── SAVE ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\nSaved → {OUTPUT}")
