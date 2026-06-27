"""Apply same fixes to Chapter 7 as were done for Ch5/Ch6"""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"
doc = Document(DOC_PATH)
paras = doc.paragraphs

# ── Find key indices ──────────────────────────────────────────────────────────
ch7_title_idx    = None   # "CHAPTER 7"
ch7_subtitle_idx = None   # "CONCLUSION AND FUTURE WORK"
ch7_first_idx    = None   # "7.1 Summary of Work"
lot_last_el      = None   # last 'table of figures' paragraph element

for i, p in enumerate(paras):
    t = p.text.strip()
    if t == "CHAPTER 7" and ch7_title_idx is None:
        ch7_title_idx = i
    if t == "CONCLUSION AND FUTURE WORK" and ch7_subtitle_idx is None:
        ch7_subtitle_idx = i
    if t == "7.1 Summary of Work" and ch7_first_idx is None:
        ch7_first_idx = i
    if p.style.name == 'table of figures':
        lot_last_el = p._element

print(f"ch7_title={ch7_title_idx}, ch7_subtitle={ch7_subtitle_idx}, ch7_first={ch7_first_idx}")

# ── 1. Fix Chapter 7 title page ───────────────────────────────────────────────
# Step A: Add pageBreakBefore to "CHAPTER 7" paragraph (so it starts on new page)
p_ch7_title = paras[ch7_title_idx]
pPr = p_ch7_title._element.get_or_add_pPr()
for ex in pPr.findall(qn('w:pageBreakBefore')):
    pPr.remove(ex)
pBrk = OxmlElement('w:pageBreakBefore')
pBrk.set(qn('w:val'), '1')
pPr.append(pBrk)

# Step B: Add pageBreakBefore to "7.1 Summary of Work" paragraph (content on new page)
p_ch7_first = paras[ch7_first_idx]
pPr2 = p_ch7_first._element.get_or_add_pPr()
for ex in pPr2.findall(qn('w:pageBreakBefore')):
    pPr2.remove(ex)
pBrk2 = OxmlElement('w:pageBreakBefore')
pBrk2.set(qn('w:val'), '1')
pPr2.append(pBrk2)

# Step C: Insert 4 blank paragraphs before "7.1 Summary of Work"
ch7_first_el = paras[ch7_first_idx]._element
for _ in range(4):
    new_p = OxmlElement('w:p')
    ch7_first_el.addprevious(new_p)

print("OK Chapter 7 title page fixed")

# ── 2. Apply Heading styles ────────────────────────────────────────────────────
paras = doc.paragraphs  # re-read after insertions
ch7_start_idx = None
for i, p in enumerate(paras):
    if p.text.strip() == "CHAPTER 7":
        ch7_start_idx = i
        break

ch7_sec = re.compile(r'^7\.(\d+)\s')
ch7_sub = re.compile(r'^7\.(\d+)\.(\d+)\s')
tbl_cap = re.compile(r'^Table 7\.\d+:')

heading2_style = doc.styles['Heading 2']
heading3_style = doc.styles['Heading 3']
caption_style  = doc.styles['Caption']

changed_h2 = 0; changed_h3 = 0; changed_cap = 0

for i, p in enumerate(paras):
    if i < ch7_start_idx:
        continue
    t = p.text.strip()
    if not t:
        continue

    if tbl_cap.match(t):
        p.style = caption_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
        changed_cap += 1

    elif ch7_sub.match(t):
        p.style = heading3_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
        changed_h3 += 1

    elif ch7_sec.match(t):
        p.style = heading2_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
        changed_h2 += 1

print(f"OK Headings: {changed_h2} Heading 2, {changed_h3} Heading 3, {changed_cap} Captions")

# ── 3. Add Chapter 7 tables to List of Tables ────────────────────────────────
paras = doc.paragraphs
# Re-find last LOT entry
for p in paras:
    if p.style.name == 'table of figures':
        lot_last_el = p._element

ch7_tables = [
    "Table 7.1 Project Objectives and Achievement Status",
]

def make_lot_entry(doc, text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'TableofFigures')
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Times New Roman')
    rFont.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFont)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    r.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_el)
    p.append(r)
    # tab + dash for page number
    r2 = OxmlElement('w:r')
    tab_el = OxmlElement('w:tab'); r2.append(tab_el)
    p.append(r2)
    r3 = OxmlElement('w:r')
    rPr3 = OxmlElement('w:rPr')
    rFont3 = OxmlElement('w:rFonts')
    rFont3.set(qn('w:ascii'), 'Times New Roman')
    rFont3.set(qn('w:hAnsi'), 'Times New Roman')
    rPr3.append(rFont3)
    sz3 = OxmlElement('w:sz'); sz3.set(qn('w:val'), '24'); rPr3.append(sz3)
    r3.append(rPr3)
    t3 = OxmlElement('w:t'); t3.text = '-'; r3.append(t3)
    p.append(r3)
    return p

insert_after = lot_last_el
for entry in ch7_tables:
    new_entry = make_lot_entry(doc, entry)
    insert_after.addnext(new_entry)
    insert_after = new_entry

print(f"OK Added {len(ch7_tables)} entries to List of Tables")

# ── 4. Save and verify ────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")

# Quick verify
doc2 = Document(DOC_PATH)
paras2 = doc2.paragraphs
print("\n=== VERIFY: Ch7 headings ===")
ch7_pat = re.compile(r'^7\.\d')
for i, p in enumerate(paras2):
    t = p.text.strip()
    if ch7_pat.match(t):
        print(f"  [{i}] {p.style.name} | '{t}'")

print("\n=== VERIFY: Ch7 title page break ===")
for i, p in enumerate(paras2):
    t = p.text.strip()
    if t in ("CHAPTER 7", "CONCLUSION AND FUTURE WORK", "7.1 Summary of Work"):
        has_pb = 'w:pageBreakBefore' in p._element.xml
        print(f"  [{i}] pb={has_pb} style='{p.style.name}' | '{t}'")
