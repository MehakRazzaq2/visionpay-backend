"""Pre-check: show all body elements (p + tbl) in order, track chapters"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx")
body = doc.element.body

def get_text(el):
    parts = []
    for t in el.iter(qn('w:t')):
        if t.text:
            parts.append(t.text)
    return ''.join(parts).strip()

def get_style(el):
    s = el.find('.//' + qn('w:pStyle'))
    return s.get(qn('w:val'), '') if s is not None else 'Normal'

chapter = 0
table_idx = 0
print("=== Body elements: chapters and tables ===")
for idx, el in enumerate(body):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag == 'p':
        text = get_text(el)
        style = get_style(el)
        m = re.match(r'^(?:CHAPTER|Chapter)\s+(\d)', text)
        if m:
            chapter = int(m.group(1))
            print(f"\n--- CHAPTER {chapter} starts at body[{idx}] ---")
        # Show Caption paragraphs
        if 'Caption' in style or 'caption' in style.lower():
            print(f"  p[{idx}] CAPTION: '{text[:80]}'")
    elif tag == 'tbl':
        table_idx += 1
        # get first cell text
        rows = el.findall('.//' + qn('w:tr'))
        first_cell_text = ''
        if rows:
            cells = rows[0].findall('.//' + qn('w:tc'))
            if cells:
                first_cell_text = get_text(cells[0])
        # check prev element
        prev = list(body)[idx-1] if idx > 0 else None
        prev_tag = prev.tag.split('}')[-1] if prev is not None and '}' in prev.tag else ''
        prev_style = ''
        if prev_tag == 'p':
            prev_style = get_style(prev)
            prev_text = get_text(prev)
        else:
            prev_text = ''
        prev_is_cap = 'Caption' in prev_style or 'caption' in prev_style.lower()
        print(f"  tbl[{idx}] TABLE#{table_idx} Ch{chapter} | first_cell='{first_cell_text[:40]}' | prev_cap={prev_is_cap} | prev='{prev_text[:50]}'")

print(f"\nTotal body elements: {len(list(body))}")
print(f"Total tables: {table_idx}")
print(f"Total doc.paragraphs: {len(doc.paragraphs)}")
print(f"Total doc.tables: {len(doc.tables)}")
