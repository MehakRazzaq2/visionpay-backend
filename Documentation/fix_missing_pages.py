"""Fix the 4 LOT entries that didn't get page numbers - direct approach"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"

# These are the 4 missing ones with their correct page numbers from Word
MISSING = {
    "Table 5.2 Hardware Components Used": 63,
    "Table 5.6 Load Cell to HX711 Wiring": 70,
    "Table 6.3 Overall Model Performance Metrics": 78,
    "Table 6.8 Comparison with Existing Billing Systems": 85,
}

doc = Document(DOC_PATH)

# Debug: show all 'table of figures' paragraphs for 5.x/6.x/7.x
print("=== Current LOT entries for Ch5/6/7 ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'table of figures':
        t = p.text
        if any(x in t for x in ['Table 5.', 'Table 6.', 'Table 7.']):
            print(f"[{i}] repr={repr(t[:60])}")

print()

fixed = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name != 'table of figures':
        continue
    raw = p.text
    for lot_name, pg in MISSING.items():
        # The name appears at the START of the text (before the tab)
        if raw.strip().startswith(lot_name) or lot_name in raw:
            print(f"Fixing [{i}]: '{lot_name}' -> page {pg}")
            # Clear all runs completely
            for run in p.runs:
                run.text = ""
            # Rewrite using runs properly
            # First clear the entire paragraph XML content (keep pPr)
            p_el = p._element
            pPr = p_el.find(qn('w:pPr'))
            # Remove all runs and other content
            for child in list(p_el):
                tag = child.tag.split('}')[-1]
                if tag != 'pPr':
                    p_el.remove(child)
            # Add run with lot_name text
            r1 = OxmlElement('w:r')
            rPr1 = OxmlElement('w:rPr')
            rf1 = OxmlElement('w:rFonts')
            rf1.set(qn('w:ascii'), 'Times New Roman')
            rf1.set(qn('w:hAnsi'), 'Times New Roman')
            rPr1.append(rf1)
            sz1 = OxmlElement('w:sz'); sz1.set(qn('w:val'), '24')
            rPr1.append(sz1)
            r1.append(rPr1)
            t1 = OxmlElement('w:t')
            t1.text = lot_name
            t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r1.append(t1)
            p_el.append(r1)
            # Add tab run
            r2 = OxmlElement('w:r')
            tab_el = OxmlElement('w:tab')
            r2.append(tab_el)
            p_el.append(r2)
            # Add page number run
            r3 = OxmlElement('w:r')
            rPr3 = OxmlElement('w:rPr')
            rf3 = OxmlElement('w:rFonts')
            rf3.set(qn('w:ascii'), 'Times New Roman')
            rf3.set(qn('w:hAnsi'), 'Times New Roman')
            rPr3.append(rf3)
            sz3 = OxmlElement('w:sz'); sz3.set(qn('w:val'), '24')
            rPr3.append(sz3)
            r3.append(rPr3)
            t3 = OxmlElement('w:t')
            t3.text = str(pg)
            r3.append(t3)
            p_el.append(r3)
            fixed += 1
            break

print(f"\nFixed {fixed} entries")
doc.save(DOC_PATH)
print("Saved.")

# Final verification
print("\n=== FINAL: All Ch5/6/7 LOT entries ===")
doc2 = Document(DOC_PATH)
for i, p in enumerate(doc2.paragraphs):
    if p.style.name == 'table of figures':
        t = p.text.strip()
        if any(x in t for x in ['Table 5.', 'Table 6.', 'Table 7.']):
            print(f"  {t}")
