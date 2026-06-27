"""
Fix FYP Report - Copy.docx directly:
1. Chapter 6 title page: add blank lines + page break before 6.1
2. Ch5 and Ch6 headings: Normal → Heading 2 / Heading 3  (for TOC)
3. Table captions in Ch5/Ch6: Normal → Caption  (for List of Tables)
4. Add Ch5/Ch6 tables to List of Tables
"""
import re, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"

doc = Document(DOC_PATH)
paras = doc.paragraphs

# ── Find key paragraph indices ────────────────────────────────────────────────
ch5_start_idx  = None   # "Chapter 5"
ch6_title_idx  = None   # "TESTING AND EVALUATION"
ch6_first_idx  = None   # "6.1 Overview"
lot_heading_idx = None  # "List Of Tables"
lot_last_idx   = None   # last 'table of figures' entry in List of Tables

for i, p in enumerate(paras):
    t = p.text.strip()
    if t == "Chapter 5" and ch5_start_idx is None:
        ch5_start_idx = i
    if t == "TESTING AND EVALUATION" and ch6_title_idx is None:
        ch6_title_idx = i
    if t == "6.1 Overview" and ch6_first_idx is None:
        ch6_first_idx = i
    if "List Of Tables" in t and lot_heading_idx is None:
        lot_heading_idx = i
    if p.style.name == 'table of figures' and lot_heading_idx is not None:
        lot_last_idx = i

print(f"ch5_start={ch5_start_idx}, ch6_title={ch6_title_idx}, ch6_first={ch6_first_idx}")
print(f"lot_heading={lot_heading_idx}, lot_last={lot_last_idx}")

# ── 1. Fix Chapter 6 Title Page ───────────────────────────────────────────────
# Chapter 5 has: title para → subtitle para → 5 blank paras → content (with pageBreakBefore)
# Chapter 6 has: title para → subtitle para → content IMMEDIATELY (no break!)
# Fix: add 4 blank paras + pageBreakBefore on 6.1

# Add pageBreakBefore to 6.1 Overview paragraph
p6_first = paras[ch6_first_idx]
pPr = p6_first._element.get_or_add_pPr()
# Remove existing pageBreakBefore if any
for existing in pPr.findall(qn('w:pageBreakBefore')):
    pPr.remove(existing)
pBrk = OxmlElement('w:pageBreakBefore')
pBrk.set(qn('w:val'), '1')
pPr.append(pBrk)

# Insert 4 blank Normal paragraphs BEFORE "6.1 Overview" element
ch6_first_el = paras[ch6_first_idx]._element
for _ in range(4):
    new_p = OxmlElement('w:p')
    ch6_first_el.addprevious(new_p)

print("OK Chapter 6 title page fixed")

# ── 2. Apply Heading Styles to Ch5 and Ch6 ────────────────────────────────────
# Re-read paras after XML changes (insertions above changed the list)
paras = doc.paragraphs

ch5_sec   = re.compile(r'^5\.(\d+)\s')   # 5.1, 5.2, ...
ch5_sub   = re.compile(r'^5\.(\d+)\.(\d+)\s')  # 5.2.1, 5.4.2, ...
ch6_sec   = re.compile(r'^6\.(\d+)\s')
ch6_sub   = re.compile(r'^6\.(\d+)\.(\d+)\s')
tbl_cap   = re.compile(r'^Table [56]\.\d+:')

# We need fresh ch5_start after re-read
ch5_start_idx = None
for i, p in enumerate(paras):
    if p.text.strip() == "Chapter 5":
        ch5_start_idx = i
        break

heading2_style = doc.styles['Heading 2']
heading3_style = doc.styles['Heading 3']
caption_style  = doc.styles['Caption']

changed_h2 = 0; changed_h3 = 0; changed_cap = 0

for i, p in enumerate(paras):
    if i < ch5_start_idx:
        continue
    t = p.text.strip()
    if not t:
        continue

    if tbl_cap.match(t):
        # Table caption → Caption style
        p.style = caption_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
            run.font.color.rgb = None   # let style control colour
        changed_cap += 1

    elif ch5_sub.match(t) or ch6_sub.match(t):
        # X.X.X heading → Heading 3
        p.style = heading3_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
        changed_h3 += 1

    elif ch5_sec.match(t) or ch6_sec.match(t):
        # X.X heading → Heading 2
        p.style = heading2_style
        for run in p.runs:
            run.font.bold = None
            run.font.size = None
        changed_h2 += 1

print(f"OK Headings: {changed_h2} Heading 2, {changed_h3} Heading 3, {changed_cap} Captions")

# ── 3. Update List of Tables (add Ch5 and Ch6 entries) ───────────────────────
# Re-read paras again
paras = doc.paragraphs

# Find last List of Tables entry
lot_heading_idx = None
lot_last_el = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if "List Of Tables" in t and lot_heading_idx is None:
        lot_heading_idx = i
    if p.style.name == 'table of figures' and lot_heading_idx is not None:
        lot_last_el = p._element

ch5_tables = [
    "Table 5.1 Development Tools and Libraries",
    "Table 5.2 Hardware Components Used",
    "Table 5.3 Dataset Summary",
    "Table 5.4 YOLOv8n Training Configuration",
    "Table 5.5 Per-Class Detection Accuracy (Selected Classes)",
    "Table 5.6 Load Cell to HX711 Wiring",
    "Table 5.7 HX711 to Arduino Wiring",
    "Table 5.8 Deployment Configuration",
]
ch6_tables = [
    "Table 6.1 Testing Types and Scope",
    "Table 6.2 API Endpoint Test Results",
    "Table 6.3 Overall Model Performance Metrics",
    "Table 6.4 Per-Class mAP50 Results",
    "Table 6.5 Integration Test Scenarios",
    "Table 6.6 System Performance Results",
    "Table 6.7 Load Cell Accuracy Test Results",
    "Table 6.8 Comparison with Existing Billing Systems",
]

def make_lot_entry(doc, text, page_str="—"):
    """Create a 'table of figures' style entry matching existing LOT entries."""
    p = OxmlElement('w:p')
    # paragraph properties — use table of figures style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'tableoffigures')  # internal style id for 'table of figures'
    pPr.append(pStyle)
    p.append(pPr)
    # run with text
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Times New Roman')
    rFont.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFont)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    r.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = f"{text}"
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_el)
    p.append(r)
    # tab + page number run
    r2 = OxmlElement('w:r')
    tab_el = OxmlElement('w:tab')
    r2.append(tab_el)
    p.append(r2)
    r3 = OxmlElement('w:r')
    rPr3 = OxmlElement('w:rPr')
    rFont3 = OxmlElement('w:rFonts')
    rFont3.set(qn('w:ascii'), 'Times New Roman')
    rFont3.set(qn('w:hAnsi'), 'Times New Roman')
    rPr3.append(rFont3)
    sz3 = OxmlElement('w:sz'); sz3.set(qn('w:val'), '24'); rPr3.append(sz3)
    r3.append(rPr3)
    t3 = OxmlElement('w:t')
    t3.text = page_str
    r3.append(t3)
    p.append(r3)
    return p

if lot_last_el is not None:
    # Insert all ch5 and ch6 table entries after the last existing LOT entry
    insert_after = lot_last_el
    for entry in ch5_tables + ch6_tables:
        new_entry = make_lot_entry(doc, entry)
        insert_after.addnext(new_entry)
        insert_after = new_entry
    print(f"OK Added {len(ch5_tables)+len(ch6_tables)} entries to List of Tables")
else:
    print("WARNING: Could not find List of Tables entries to append to")

# ── 4. Save ───────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\nDONE Document saved: {DOC_PATH}")
print("\nIMPORTANT: Open the document in Word and:")
print("  1. Right-click on the TABLE OF CONTENTS → Update Field → Update entire table")
print("  2. Right-click on the LIST OF FIGURES → Update Field → Update entire table")
print("  3. Update page numbers in List of Tables manually if needed")
