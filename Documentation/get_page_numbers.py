"""
Use Word COM to get actual page numbers for each table caption,
then update the List of Tables entries in the document.
"""
import win32com.client
import re, os

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"
wdActiveEndPageNumber = 3  # Word constant

# ── Step 1: Open in Word invisibly and get page numbers ──────────────────────
word = win32com.client.Dispatch("Word.Application")
word.Visible = False

try:
    doc_com = word.Documents.Open(DOC_PATH)
    doc_com.Repaginate()  # force Word to recalculate pagination

    # Tables we need page numbers for
    targets = [
        "Table 5.1: Development Tools and Libraries",
        "Table 5.2: Hardware Components Used",
        "Table 5.3: Dataset Summary",
        "Table 5.4: YOLOv8n Training Configuration",
        "Table 5.5: Per-Class Detection Accuracy (Selected Classes)",
        "Table 5.6: Load Cell to HX711 Wiring",
        "Table 5.7: HX711 to Arduino Wiring",
        "Table 5.8: Deployment Configuration",
        "Table 6.1: Testing Types and Scope",
        "Table 6.2: API Endpoint Test Results",
        "Table 6.3: Overall Model Performance Metrics",
        "Table 6.4: Per-Class mAP50 Results",
        "Table 6.5: Integration Test Scenarios",
        "Table 6.6: System Performance Results",
        "Table 6.7: Load Cell Accuracy Test Results",
        "Table 6.8: Comparison with Existing Billing Systems",
        "Table 7.1: Project Objectives and Achievement Status",
    ]

    page_map = {}  # "Table X.Y: ..." -> page number

    for para in doc_com.Paragraphs:
        text = para.Range.Text.strip()
        if text in targets:
            pg = para.Range.Information(wdActiveEndPageNumber)
            page_map[text] = pg
            print(f"  Found: '{text}' -> page {pg}")

    doc_com.Close(False)
finally:
    word.Quit()

print(f"\nPage numbers found: {len(page_map)}/{len(targets)}")

# ── Step 2: Update List of Tables entries in the docx ────────────────────────
from docx import Document

doc = Document(DOC_PATH)
paras = doc.paragraphs

# LOT entries format: "Table X.Y Name\ttab-PageNum"
# We match by looking at the text (without colon) in LOT entries
# Caption: "Table 5.1: Development Tools and Libraries"
# LOT entry: "Table 5.1 Development Tools and Libraries\t—"

def caption_to_lot_key(caption_text):
    """Convert 'Table 5.1: Dev...' to 'Table 5.1 Dev...' for matching"""
    return caption_text.replace(":", "", 1)  # remove first colon only

lot_entry_re = re.compile(r'^(Table [567]\.\d+ .+?)\t')

updated = 0
for p in paras:
    if p.style.name != 'table of figures':
        continue
    text = p.text
    m = lot_entry_re.match(text)
    if not m:
        continue
    lot_key = m.group(1)  # e.g. "Table 5.1 Development Tools and Libraries"

    # Find matching caption
    for caption, pg in page_map.items():
        if caption_to_lot_key(caption).startswith(lot_key[:20]):
            # Clear existing runs and rewrite with correct page number
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = lot_key + "\t" + str(pg)
            else:
                p.add_run(lot_key + "\t" + str(pg))
            updated += 1
            print(f"  Updated LOT: '{lot_key}' -> page {pg}")
            break

print(f"\nUpdated {updated} LOT entries")

doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")
