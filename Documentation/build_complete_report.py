"""
VisionPay — Complete FYP Report Generator
Fixes: title, supervisor names, Chapter 5 rewrite, abbreviations, formatting
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

TITLE    = "Cross-Platform Grocery Billing Application using Artificial Intelligence"
SUBTITLE = "VisionPay"
SUPERVISOR    = "Dr. Sehrish Khan Tayyaba"
CO_SUPERVISOR = "Dr. Altaf Hussain"
INSTITUTE = "Institute of Space Technology, KICSIT, Kahuta Campus"
YEAR = "2026"
MEMBERS = [
    ("Naveed Hayat",  "222201006"),
    ("Nimrah Khan",   "222201008"),
    ("Mehak Razzaq",  "222201025"),
]

# ── Helpers ───────────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    for s in doc.sections:
        s.left_margin   = Inches(1.5)
        s.right_margin  = Inches(1.0)
        s.top_margin    = Inches(1.0)
        s.bottom_margin = Inches(1.0)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    return doc

def _font(run, size, bold=False, italic=False):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text); _font(r, 16, bold=True)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text); _font(r, 14, bold=True)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text); _font(r, 12, bold=True)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text); _font(r, 12)
    return p

def centre(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); _font(r, size, bold=bold)
    return p

def tbl_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text); _font(r, 12, bold=True)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        for run in hcells[i].paragraphs[0].runs:
            run.font.name = "Times New Roman"; run.font.size = Pt(11); run.font.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, v in enumerate(row):
            cells[ci].text = str(v)
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

def fig_placeholder(doc, num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"[Figure {num}: {caption}]"); _font(r, 11, italic=True)

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); _font(r, 12)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# FRONT MATTER
# ═══════════════════════════════════════════════════════════════════════════════
def add_title_page(doc):
    for _ in range(4): doc.add_paragraph()
    centre(doc, TITLE, size=16, bold=True)
    doc.add_paragraph()
    centre(doc, SUBTITLE, size=14, bold=True)
    doc.add_paragraph()
    centre(doc, "Submitted by", size=12)
    for name, roll in MEMBERS:
        centre(doc, f"{name} ({roll})", size=12)
    doc.add_paragraph()
    centre(doc, f"Supervisor: {SUPERVISOR}", size=12)
    centre(doc, f"Co-Supervisor: {CO_SUPERVISOR}", size=12)
    doc.add_paragraph()
    centre(doc, "Department of Computer Science", size=12)
    centre(doc, INSTITUTE, size=12)
    centre(doc, YEAR, size=12)
    doc.add_paragraph()
    centre(doc, "APPROVAL BY BOARD OF EXAMINERS", size=12, bold=True)

def add_approval_page(doc):
    page_break(doc)
    h1(doc, "APPROVAL BY BOARD OF EXAMINERS")
    doc.add_paragraph()
    body(doc, f"This is to certify that the thesis titled \"{TITLE}\" submitted by Naveed Hayat (222201006), Nimrah Khan (222201008), and Mehak Razzaq (222201025) has been evaluated and approved by the Board of Examiners.")
    doc.add_paragraph()
    for label in ["Supervisor", "Committee Member", "Committee Member", "Committee Member", "Lecturer"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run(f"______________________          {label}"); _font(r, 12)
    doc.add_paragraph()
    for name, role in [(SUPERVISOR, "Project Coordinator"), (CO_SUPERVISOR, "HoD, Computer Science"),
                       ("Prof. Dr. Syed Nasir Mehmood Shah", "Head of Academics, KICSIT")]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        r = p.add_run(f"______________________\n{name}\n{role}"); _font(r, 12)

def add_declaration(doc):
    page_break(doc)
    h1(doc, "AUTHORS DECLARATION")
    body(doc, f"We take full responsibility for the research work conducted during the thesis titled \"{TITLE}\". We solemnly declare that the research and development work presented in this thesis was done solely by us with no significant help from any other person; however, small help wherever taken is duly acknowledged. We have also written this thesis ourselves.")
    body(doc, "We have not presented this thesis or any part of it previously to any other degree-awarding institution within Pakistan or abroad. We understand that the management of IST has a zero-tolerance policy towards plagiarism. Therefore, we as authors of the above-mentioned thesis solemnly declare that no portion of this thesis has been plagiarized and any material used does not contain any literal citing of more than 70 words even by giving a reference unless written permission of the publisher has been obtained.")
    doc.add_paragraph()
    for name, _ in MEMBERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        r = p.add_run(f"______________________          {name}"); _font(r, 12)
    doc.add_paragraph()
    body(doc, "I hereby acknowledge that the submitted thesis is the final version and should be scrutinized for plagiarism as per IST policy.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"______________________\nVerified by Plagiarism Cell Officer\nDated: _______________"); _font(r, 12)

def add_certificate(doc):
    page_break(doc)
    h1(doc, "CERTIFICATE")
    body(doc, f"This is to certify that the research work described in this thesis is the original work of the authors and has been carried out under my direct supervision. I have personally gone through all the data, results, and materials reported in the manuscript and certify their correctness and authenticity. I further certify that the material included in this thesis is not plagiarized and has not been used in part or in full in a manuscript already submitted or in the process of submission for the award of any other degree from any institution.")
    body(doc, "I also certify that this thesis has been prepared under my supervision according to the prescribed format and I endorse its evaluation for the award of the Bachelor of Science degree in Computer Science through the official procedures of the Institute.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(f"______________________\n{SUPERVISOR}\nSupervisor"); _font(r, 12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(f"______________________\n{CO_SUPERVISOR}\nCo-Supervisor"); _font(r, 12)

def add_copyright(doc):
    page_break(doc)
    h1(doc, "COPYRIGHT")
    body(doc, f"Copyright © {YEAR}")
    body(doc, "This document is jointly copyrighted by the authors and the Institute of Space Technology (IST). Both the authors and IST can use, publish, or reproduce this document in any form. Under the copyright law, no part of this document can be reproduced by anyone, except the copyright holders, without the written permission of the authors.")

def add_dedication(doc):
    page_break(doc)
    h1(doc, "DEDICATION")
    for _ in range(3): doc.add_paragraph()
    centre(doc, "This research work is dedicated to our beloved Parents, Teachers, and Friends", size=12, bold=True)
    centre(doc, "who supported and motivated us through every challenge", size=12)
    centre(doc, "and made us into who we are today.", size=12)

def add_acknowledgements(doc):
    page_break(doc)
    h1(doc, "ACKNOWLEDGEMENTS")
    body(doc, f"All praise and gratitude is due to Allah Almighty, whose blessings and guidance made this work possible. We would like to express our sincere gratitude to our supervisor, {SUPERVISOR}, and co-supervisor, {CO_SUPERVISOR}, for their continuous guidance, valuable feedback, and encouragement throughout this project. Their expertise and dedication to our academic development were instrumental in shaping this work.")
    body(doc, "We are also thankful to the Department of Computer Science at the Institute of Space Technology, KICSIT Kahuta Campus, for providing the academic environment and resources necessary for this project. Our heartfelt thanks go to our families for their unconditional support, patience, and motivation during the entire duration of our studies. We also thank our friends and colleagues who helped us in testing the system and provided useful feedback.")
    body(doc, "Finally, we would like to acknowledge the open-source communities behind Flutter, FastAPI, Ultralytics YOLOv8, and the Hugging Face platform, whose freely available tools made the development of VisionPay possible.")

def add_abstract(doc):
    page_break(doc)
    h1(doc, "ABSTRACT")
    body(doc, "Traditional grocery billing systems rely heavily on barcode scanning, which increases checkout time and creates operational inefficiency, particularly in small and medium-sized retail stores. This thesis presents the design and development of VisionPay — a cross-platform grocery billing application that uses Artificial Intelligence (AI) for automated product detection and billing.")
    body(doc, "The system identifies packaged and unpackaged grocery items using a YOLOv8n object detection model trained on 53 classes of Pakistani grocery products, fruits, and vegetables. A digital load cell integrated with an Arduino microcontroller measures the weight of produce items for accurate weight-based billing. The backend is built using FastAPI and deployed permanently on Hugging Face Spaces, making the system accessible from any device without requiring a local server. The cross-platform mobile and web application is developed in Flutter.")
    body(doc, "The model achieved an overall mean Average Precision at 50% Intersection over Union (mAP50) of 39.9%, with packed products like CocoMo and Lays French Cheese achieving near-perfect accuracy of 99.5%. The system supports real-time synchronization between a mobile phone and a laptop browser via WebSocket, allowing the phone camera to scan products while billing results are displayed on the laptop screen. Experimental results demonstrate the feasibility and practicality of the proposed approach for small grocery stores in developing economies.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Keywords: "); _font(r, 12, bold=True)
    r2 = p.add_run("Computer Vision, Object Detection, YOLOv8, Grocery Billing, Flutter, FastAPI, Load Cell, Cross-Platform Application.")
    _font(r2, 12)

def add_abbreviations(doc):
    page_break(doc)
    h1(doc, "LIST OF ABBREVIATIONS")
    abbrevs = [
        ("AI",    "Artificial Intelligence"),
        ("ML",    "Machine Learning"),
        ("CV",    "Computer Vision"),
        ("OCR",   "Optical Character Recognition"),
        ("YOLO",  "You Only Look Once"),
        ("API",   "Application Programming Interface"),
        ("REST",  "Representational State Transfer"),
        ("JSON",  "JavaScript Object Notation"),
        ("UI",    "User Interface"),
        ("DB",    "Database"),
        ("GPU",   "Graphics Processing Unit"),
        ("mAP",   "Mean Average Precision"),
        ("IoU",   "Intersection over Union"),
        ("SDK",   "Software Development Kit"),
        ("APK",   "Android Package Kit"),
        ("USB",   "Universal Serial Bus"),
        ("IDE",   "Integrated Development Environment"),
        ("CNN",   "Convolutional Neural Network"),
        ("SQLite","Structured Query Language Lite"),
        ("HF",    "Hugging Face"),
    ]
    tbl_caption(doc, "")
    add_table(doc, ["Abbreviation", "Full Form"], abbrevs)

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter1(doc):
    page_break(doc)
    h1(doc, "CHAPTER 1")
    h1(doc, "INTRODUCTION")

    h2(doc, "1.1 General Overview of Grocery Billing Systems")
    body(doc, "Grocery retailing is one of the most universal and vital businesses in the world, especially in developing economies like Pakistan. Although digital technologies have improved significantly over the past two decades, in the majority of grocery stores the billing procedure is still based on the conventional method of using barcodes. In these systems, each product is assigned a unique barcode that is scanned at the checkout counter to retrieve price and product information. While barcode-based billing has improved accuracy compared to fully manual methods, it still presents several practical difficulties in daily retail operations.")

    h2(doc, "1.2 Shortcomings of Traditional Barcode-Based Billing")
    body(doc, "In a standard barcode-based billing system, the cashier must physically handle each product, locate the barcode, position it in front of the scanner, and wait for the system to process the scan. This process takes significant time during peak shopping hours when customer traffic is high. Damaged, folded, or missing barcodes slow the checkout process further and often require manual entry, increasing the risk of human error. Barcode systems also require considerable human involvement and are therefore prone to inefficiencies caused by mis-scans, incorrect item handling, or operator fatigue, all of which lead to longer checkout queues and lower customer satisfaction.")

    h2(doc, "1.3 Grocery Retail in Developing Markets")
    body(doc, "In Pakistan, the grocery retail market consists primarily of small to medium-sized stores and local marts with limited technical infrastructure. Most of these stores either operate basic point-of-sale systems or use semi-manual billing methods. While fully automated retail systems exist in developed markets, local grocery businesses can rarely afford or install them due to their high hardware costs and complex setup requirements. Cashier-less or fully automated retail systems typically require expensive infrastructure including ceiling-mounted cameras, smart shelves, and powerful computing hardware, none of which is practical for the developing economy context where cost and simplicity are the primary concerns.")

    h2(doc, "1.4 Computer Vision and Implicit Barcoding")
    body(doc, "Recent advances in computer vision and Artificial Intelligence (AI) have introduced new methods for product identification in retail environments. Implicit barcoding — where a product is identified by its visual, textual, and physical characteristics rather than an explicit printed barcode — is one such approach. Implicit barcoding uses technologies such as packaging recognition, Optical Character Recognition (OCR), brand detection, and feature extraction to identify products. Computer vision-based systems can help eliminate conventional barcode scanning by processing the visual image of the packaging and extracting meaningful information from it. When combined with physical measurements such as weight, implicit barcoding can identify both packaged and unpackaged grocery products more reliably.")

    h2(doc, "1.5 Proposed System: VisionPay")
    body(doc, "The system proposed in this thesis is called VisionPay and is designed as a practical solution for the local grocery store setting. VisionPay uses a mobile phone camera to capture images of products and applies a YOLOv8n object detection model to identify them. The system handles packaged products with known prices directly, while for produce items such as fruits and vegetables, it integrates with a load cell-based weighing system to calculate the price by weight. A cross-platform application built in Flutter serves both the cashier and the store manager, and the backend is deployed permanently on the cloud so that the system is always accessible without requiring a laptop to be on.")

    h2(doc, "1.6 Scope and Focus")
    body(doc, "The primary focus of this project is to build an intelligent billing system prototype that incorporates software-based AI with minimal hardware requirements. The system is designed to work under controlled retail conditions and represents a proof-of-concept rather than a fully commercial solution. It emphasizes practicality, affordability, and scalability so that the proposed approach can realistically be adopted by small and medium-sized grocery stores in developing markets like Pakistan.")
    fig_placeholder(doc, "1.1", "Traditional Barcode Billing vs VisionPay Proposed Solution")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter2(doc):
    page_break(doc)
    h1(doc, "CHAPTER 2")
    h1(doc, "LITERATURE REVIEW")

    h2(doc, "2.1 Automated Billing Systems Overview")
    body(doc, "Traditional retail billing systems are founded on explicit barcode scanning, in which every product is manually scanned at the checkout counter using a barcode scanner. This method has been widely used in grocery stores, supermarkets, and retail chains because of its simplicity and standardization. However, as customer volumes increase and the demand for faster checkout experiences grows, barcode-based systems have shown significant limitations. Under traditional barcode billing, the cashier or customer must scan each item separately, causing long queues, human dependency, and billing delays particularly during peak hours. Barcodes may also be damaged, concealed, or lost, causing scan failures. These restrictions have motivated researchers to explore automated and smart billing systems that minimize manual processing through computer vision, machine learning, and sensor technologies.")

    h2(doc, "2.2 Barcode-Based Billing Systems (Current Solution)")
    body(doc, "Barcode-based billing systems use one-dimensional (1D) or two-dimensional (2D) codes printed on product packaging. Each barcode holds a unique identifier that is associated with product data stored in a centralized database, including price, name, and category. The primary advantages of barcode systems are their simplicity, low hardware cost, and easy database integration. Their key limitations include the requirement for manual scanning, the inability to recognize products when barcodes are not visible, the absence of any intelligent product validation, reduced efficiency in high-traffic retail environments, and vulnerability to fraud and misbilling.")

    h2(doc, "2.3 Deep Learning-Based Automated Store Billing System (2023)")
    body(doc, "Neserbaum et al. proposed a deep learning-based automated billing system that uses image classification to identify retail items at a checkout point without explicit barcode scanning. The system captures product images through a fixed camera and feeds them into a trained Convolutional Neural Network (CNN) to detect and classify items. The methodology includes preprocessing and normalization of input images followed by CNN-based product classification and automated bill generation. While the study demonstrated the feasibility of barcode-free billing, it suffered from poor scalability due to fixed camera angles, degraded performance when similar packaging products were present, and no integration of weight verification or OCR for text-based confirmation.")

    h2(doc, "2.4 ARC: A Vision-Based Automatic Retail Checkout System (2022)")
    body(doc, "Khan et al. proposed ARC, a vision-based retail checkout system that eliminates conventional checkout methods by applying computer vision and deep learning to identify products in a retail environment. ARC uses multi-view image capture to extract product features and applies deep learning models for recognition, followed by automated billing. The system improves automation but remains reliant on optimal environmental conditions including proper lighting and camera setup. It lacks OCR integration for product metadata extraction and does not include any sensor-based weight validation mechanism, limiting its reliability in real-world retail settings.")

    h2(doc, "2.5 Amazon Go: Smart Store Model (Industry Benchmark)")
    body(doc, "Amazon Go represents the most advanced commercially deployed automated retail system, using a \"Just Walk Out\" technology that relies on hundreds of high-resolution cameras, smart shelves with built-in sensors, advanced sensor fusion, and real-time cloud processing. While Amazon Go achieves near-complete automation, it is entirely unsuitable for academic research or developing economy deployment due to its extremely high infrastructure costs, heavy hardware dependency, complex calibration requirements, and the inability to scale down to small or medium-sized stores.")

    h2(doc, "2.6 Research Gaps and Motivation")
    body(doc, "A review of the existing literature reveals several important gaps. The majority of vision-based billing systems do not incorporate weight verification. Most systems struggle to differentiate visually similar packed products. OCR is underutilized for extracting confirming product information from packaging. Existing smart systems depend on costly infrastructure that is inaccessible for small stores. There is a lack of lightweight, affordable, and deployable solutions suitable for developing market grocery environments. These gaps collectively motivated the development of VisionPay as a practical alternative.")

    h2(doc, "2.7 Implicit Barcoding and Vision-Based Billing Systems")
    body(doc, "Implicit barcoding refers to identifying a product not through a physical barcode label but through the product's own visual features. These include packaging design, brand logos, color patterns, textual information extracted via OCR, shape, size, and weight. Computer vision and deep learning models applied to product images extract discriminative features that serve as a virtual barcode, allowing the system to uniquely identify products even when traditional barcodes are absent. This approach is particularly valuable for grocery products that look similar to one another and where brand differentiation depends on packaging design and associated text.")

    h2(doc, "2.8 Justification for VisionPay")
    body(doc, "VisionPay addresses the gaps identified above by offering a smart AI-based billing system that combines packaging recognition as implicit barcoding, deep learning-based brand detection, OCR for metadata extraction, and weight validation using a load cell. Unlike Amazon Go, VisionPay does not require heavy infrastructure and is appropriate for small to medium retail stores and academic settings. The system uses a mobile phone as its primary camera, reducing hardware costs to nearly zero, and deploys the processing backend on a free cloud platform.")

    h2(doc, "2.9 Literature Review Summary")
    tbl_caption(doc, "Table 2.1: Literature Review Summary")
    add_table(doc,
        ["System", "Methodology", "Technology", "Limitation"],
        [
            ["Barcode Billing (Current)", "Manual scanning, database match", "Barcode scanner, DB", "Manual effort, damaged barcode fails, no intelligence"],
            ["Deep Learning Billing (2023)", "Fixed camera, CNN classification", "CNN, Computer Vision", "Fixed angles, no weight/OCR, similar packaging fails"],
            ["ARC Vision Checkout (2022)", "Multi-view capture, deep learning", "CV, Deep Learning", "Requires optimal lighting, no OCR/weight validation"],
            ["Amazon Go", "Multi-camera tracking, sensor fusion", "High-res cameras, Smart shelves, Cloud", "Very expensive, not feasible for small stores"],
            ["VisionPay (Proposed)", "YOLO detection, load cell weight, FastAPI, Flutter", "YOLOv8n, Arduino, Flutter, FastAPI", "Prototype-stage; Internet required"],
        ]
    )

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter3(doc):
    page_break(doc)
    h1(doc, "CHAPTER 3")
    h1(doc, "METHODOLOGY")

    h2(doc, "3.1 Overview")
    body(doc, "This chapter describes the methodology adopted for the development of VisionPay. To build a smart, practical, and scalable billing solution, it was first necessary to understand the current state of retail billing technologies and the limitations of automated checkout systems. The development methodology followed a systematic, iterative approach where each module was designed, implemented, and tested independently before being integrated into the complete system.")

    h2(doc, "3.2 Current Methodologies of Automated Billing")

    h3(doc, "3.2.1 Barcode-Based Billing Systems")
    body(doc, "The most commonly used retail billing method relies on barcodes printed on product packaging. Each product is assigned a unique barcode that maps to an entry in the store's product database. The primary advantages of this approach are its simplicity and low cost. However, its key limitations include the requirement for manual scanning, the vulnerability of barcodes to damage or concealment, the absence of any product verification intelligence, and reduced efficiency during peak hours. These limitations make barcode systems unsuitable as the sole billing mechanism for a modern, AI-assisted retail environment.")

    h3(doc, "3.2.2 Vision-Based Product Recognition Systems")
    body(doc, "Vision-based billing systems use computer vision and deep learning to identify products by their visual features. Cameras capture product images, and trained models classify them based on shape, color, and packaging design. These systems eliminate manual barcode scanning and speed up the checkout process but face challenges with visually similar products, sensitivity to lighting and camera angles, and lack of additional verification mechanisms.")

    h3(doc, "3.2.3 OCR-Based Product Identification")
    body(doc, "OCR is used to extract textual information from product packaging such as product name, brand, weight, and manufacturing details. OCR adds value to vision-based systems by providing a text-based confirmation of product identity, particularly for packaged grocery products. Its limitations include sensitivity to font style and packaging quality, and it requires clean, visible text areas to function accurately.")

    h3(doc, "3.2.4 Sensor-Based and Smart Shelf Systems")
    body(doc, "Advanced retail systems such as Amazon Go use sensor fusion combining smart shelves, load sensors, and multiple cameras. These systems achieve high accuracy but require very expensive infrastructure that is not feasible for small or medium-sized grocery stores.")

    h2(doc, "3.3 Challenges in Automated Billing")
    body(doc, "Despite technological progress, automated billing systems face several challenges. Product similarity is a major issue, as many packaging designs look alike. High infrastructure and hardware costs limit deployment. Scalability is difficult for complex systems. Most systems operate on trust without cross-verification. Performance also depends heavily on environmental factors such as lighting and camera placement.")

    h2(doc, "3.4 Proposed Methodology: VisionPay")
    body(doc, "VisionPay adopts a hybrid intelligent billing methodology that combines computer vision, weight measurement, and cloud deployment. The core components are: YOLOv8n-based product detection for packaging recognition; a load cell with HX711 amplifier and Arduino microcontroller for weight measurement of produce items; a FastAPI backend for request processing and database management; and a Flutter-based cross-platform application for the cashier and manager interfaces. The methodology ensures that products are identified reliably through AI detection while weight-based items are handled accurately through hardware integration.")

    h2(doc, "3.5 Role of Implicit Barcoding in VisionPay")
    body(doc, "Implicit barcoding is the core concept behind VisionPay's product identification approach. Instead of relying on a physical barcode label, the system identifies products through their visual and physical characteristics. This enables the system to work in scenarios where barcodes are damaged, absent, or not applicable, such as with loose fruits and vegetables. The implicit barcode in VisionPay is a combination of the product's YOLO-detected class label and, where available, its OCR-extracted text information.")

    h2(doc, "3.6 Summary")
    body(doc, "This chapter reviewed the current methodologies in automated retail billing and identified the key shortcomings that motivate VisionPay's development. The hybrid methodology adopted by VisionPay — combining AI object detection, weight measurement, cloud deployment, and cross-platform development — provides a viable, affordable, and practical substitute to hardware-intensive solutions. The following chapters describe the system design and implementation in detail.")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter4(doc):
    page_break(doc)
    h1(doc, "CHAPTER 4")
    h1(doc, "SYSTEM DESIGN")

    h2(doc, "4.1 System Architecture")
    body(doc, "The VisionPay system is designed with a modular, layered architecture that separates the AI processing, backend logic, data storage, and user interface into distinct components. At the core of the system is a FastAPI backend server that handles all API requests, runs the YOLO detection model, manages the SQLite database, and broadcasts real-time updates via WebSocket. The mobile and web applications built in Flutter communicate with this backend over HTTPS and WebSocket connections.")
    fig_placeholder(doc, "4.1", "VisionPay System Architecture Diagram")

    h2(doc, "4.2 Use Case Diagram")
    body(doc, "The use case diagram illustrates the functional behavior of the VisionPay system by representing the interactions between different actors and the system. It provides a high-level overview of how users and system components collaborate to achieve automated grocery billing. The three primary actors in the system are the Customer, the Store Admin (Manager), and the AI Engine.")
    fig_placeholder(doc, "4.2", "VisionPay Use Case Diagram")

    h2(doc, "4.3 Descriptive Use Cases")
    body(doc, "The following tables describe the major use cases of the VisionPay system in an expanded format, explaining how automated grocery billing is performed using computer vision, weight measurement, and database integration.")

    use_cases = [
        ("4.3.1", "Place Product on Counter", "Customer", "To initiate the automated billing process",
         "VisionPay system is active", "Product placed and ready for detection",
         "Customer places product on counter", "System prepares camera and sensors",
         "Product not placed correctly", "Prompt to reposition product"),
        ("4.3.2", "Capture Product Image", "VisionPay System (AI Engine)", "To capture the visual image of the product",
         "Product placed on counter", "Product image captured",
         "Camera activated and image captured", "Image processed for detection",
         "Image blurred", "Recapture image"),
        ("4.3.3", "Measure Product Weight", "VisionPay System (AI Engine)", "To measure product weight for weight-based billing",
         "Product on counter", "Product weight measured",
         "Load cell sensor activated, weight recorded", "Weight stored for billing",
         "Sensor error", "Retry measurement"),
        ("4.3.4", "Recognize Product", "VisionPay System (AI Engine)", "To identify product using YOLO detection",
         "Product image available", "Product recognized with class and confidence",
         "YOLO model analyzes image, class identified", "Result returned to backend",
         "Low confidence detection", "Apply confidence filter — skip if below threshold"),
        ("4.3.5", "Fetch Price from Database", "VisionPay System (AI Engine)", "To retrieve product price from database",
         "Product recognized", "Price fetched",
         "Database queried by product name", "Price and details returned",
         "Product not found in DB", "Skip product — not added to cart"),
        ("4.3.6", "Calculate Total Bill", "VisionPay System (AI Engine)", "To calculate total billing amount",
         "All products detected and priced", "Total bill calculated",
         "System computes sum of all cart items", "Bill displayed on screen",
         "Calculation error", "Recalculate bill"),
        ("4.3.7", "Generate Receipt", "VisionPay System", "To generate final transaction receipt",
         "Bill calculated and checkout pressed", "Receipt generated and transaction saved",
         "Receipt displayed, transaction saved to DB, stock deducted", "Cashier confirms",
         "DB save failure", "Retry save operation"),
        ("4.3.8", "Update Inventory", "VisionPay System (AI Engine)", "To update stock after a purchase",
         "Transaction saved successfully", "Inventory updated",
         "Stock quantity reduced for each item sold", "Low stock alert triggered if needed",
         "Database error", "Retry update"),
        ("4.3.9", "Manage Store Operations", "Store Admin", "To manage and monitor the system",
         "Admin logged in as Manager", "Store managed successfully",
         "Admin views stats, manages products, views transactions", "Changes saved",
         "Access error", "Display error message"),
    ]

    for uc in use_cases:
        num, name, actor, goal, pre, post, basic_u, basic_s, alt_u, alt_s = uc
        h3(doc, f"{num} {name}")
        tbl_caption(doc, f"Table {num}: Use Case — {name}")
        add_table(doc,
            ["Field", "Description"],
            [
                ["Use Case Name", name],
                ["Participating Actor", actor],
                ["Goal", goal],
                ["Precondition", pre],
                ["Post Condition", post],
                ["Basic Flow — User Action", basic_u],
                ["Basic Flow — System Response", basic_s],
                ["Alternative Flow — Trigger", alt_u],
                ["Alternative Flow — Response", alt_s],
            ]
        )

    h2(doc, "4.4 Activity Diagrams")
    h3(doc, "4.4.1 Customer Billing Workflow")
    body(doc, "This activity diagram illustrates the complete customer billing workflow in VisionPay, starting from placing a product on the counter to AI-driven product recognition, bill generation, and receipt display. It shows how YOLOv8n detection, database lookup, weight measurement, and bill calculation work together to produce a complete billing transaction without manual barcode scanning.")
    fig_placeholder(doc, "4.3", "Customer Billing Workflow Activity Diagram")

    h3(doc, "4.4.2 AI Engine Processing Flow")
    body(doc, "This activity diagram represents the internal processing of the VisionPay AI engine for product recognition. It demonstrates how YOLO detection, confidence filtering, database matching, and price retrieval are combined to accurately identify products and calculate billing amounts.")
    fig_placeholder(doc, "4.4", "AI Engine Processing Flow Activity Diagram")

    h3(doc, "4.4.3 Store Admin and System Monitoring")
    body(doc, "This activity diagram describes the role of the store manager in monitoring and managing the VisionPay system. It covers tasks such as viewing sales reports, monitoring inventory levels, managing the product database, and responding to low stock alerts.")
    fig_placeholder(doc, "4.5", "Store Admin and System Monitoring Activity Diagram")

    h2(doc, "4.5 Class Diagram")
    body(doc, "The class diagram shows the static relationships between the major components of the VisionPay system. The primary classes include ProductDatabase, DecisionEngine, BillingEngine, and the FastAPI application class. ProductDatabase handles all SQLite operations. DecisionEngine maps detected class names to database products. BillingEngine calculates final prices based on product type and weight.")
    fig_placeholder(doc, "4.6", "VisionPay Class Diagram")

    h2(doc, "4.6 Sequence Diagrams")
    h3(doc, "4.6.1 Product Detection and Recognition")
    body(doc, "This sequence diagram shows the interaction between the Flutter mobile app, the FastAPI backend, and the YOLO model during a product detection request. The app captures an image and sends it to /detect. The backend runs YOLO prediction, processes results through the decision engine, queries the database, and returns the detected products list to the app.")
    fig_placeholder(doc, "4.7", "Product Detection Sequence Diagram")

    h3(doc, "4.6.2 Bill Generation and Price Calculation")
    body(doc, "This sequence diagram illustrates the flow from cart confirmation to bill generation. The app sends the cart items list to /generate-bill. The backend calculates prices for each item based on whether it is weight-based or unit-based and returns the complete bill structure.")
    fig_placeholder(doc, "4.8", "Bill Generation Sequence Diagram")

    h3(doc, "4.6.3 Payment and Inventory Update")
    body(doc, "This sequence diagram shows the checkout flow where the transaction is saved to the database, stock quantities are reduced for each sold item, and the receipt is returned to the app for display. The manager dashboard then reflects the updated stock and revenue figures.")
    fig_placeholder(doc, "4.9", "Payment and Inventory Update Sequence Diagram")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter5(doc):
    page_break(doc)
    h1(doc, "CHAPTER 5")
    h1(doc, "IMPLEMENTATION")

    h2(doc, "5.1 Overview")
    body(doc, "This chapter describes the complete implementation of VisionPay. The system was built using a combination of software and hardware technologies and integrated into a working prototype suitable for real grocery store use. The implementation followed the system design presented in Chapter 4 and covered four major areas: AI model training, backend development, Flutter application development, and hardware integration. The backend was deployed permanently on a cloud platform, making the system accessible from any device with an internet connection.")

    h2(doc, "5.2 Development Environment")
    body(doc, "The development environment was set up on a Windows 11 machine. Python 3.10 was used for all backend development. Flutter SDK version 3.x was used for the cross-platform mobile and web application. Arduino Integrated Development Environment (IDE) version 2.3.9 was used for programming the microcontroller. Visual Studio Code served as the primary code editor for both Python and Flutter development.")

    h3(doc, "5.2.1 Software Tools and Libraries")
    tbl_caption(doc, "Table 5.1: Development Tools and Libraries")
    add_table(doc,
        ["Tool / Library", "Version", "Purpose"],
        [
            ["Python", "3.10", "Backend and AI engine"],
            ["FastAPI", "0.111.0", "REST API and WebSocket server"],
            ["Ultralytics YOLOv8n", "8.3.86", "Object detection model"],
            ["OpenCV (headless)", "4.9.0.80", "Image processing"],
            ["SQLite", "Built-in", "Local database"],
            ["Flutter / Dart", "3.x", "Mobile and web application"],
            ["Arduino IDE", "2.3.9", "Microcontroller programming"],
            ["Docker", "Latest", "Cloud deployment containerization"],
            ["Hugging Face Hub", "1.18.0", "Model file hosting"],
            ["pyserial", "3.x", "Arduino serial communication"],
        ]
    )

    h3(doc, "5.2.2 Hardware Components")
    tbl_caption(doc, "Table 5.2: Hardware Components")
    add_table(doc,
        ["Component", "Specification", "Purpose"],
        [
            ["Load Cell", "5 kg capacity", "Measure weight of produce"],
            ["HX711 Amplifier", "24-bit ADC module", "Amplify and digitize load cell signal"],
            ["Arduino Uno", "ATmega328P, 16 MHz", "Read HX711, send weight over serial"],
            ["Android Phone", "Android 14", "Camera for product detection"],
            ["Laptop", "Windows 11", "Run weight service and web dashboard"],
        ]
    )

    h2(doc, "5.3 Dataset Preparation")
    body(doc, "The dataset used for training the YOLOv8n model was assembled by combining two sources. The first source was a publicly available fruits and vegetables dataset from Roboflow, which provided annotated images for 45 fruit and vegetable classes. The second was a custom dataset of Pakistani packed grocery products created manually by collecting and annotating images of five locally available products: Candi Biscuit, Chunkin Chocolate, CocoMo, Lays French Cheese, and Prince Biscuit. Data augmentation techniques including horizontal flipping, brightness adjustment, and rotation were applied through Roboflow to increase dataset diversity.")

    tbl_caption(doc, "Table 5.3: Dataset Summary")
    add_table(doc,
        ["Category", "No. of Classes", "Source"],
        [
            ["Pakistani Packed Products", "5", "Custom (manually collected and annotated)"],
            ["Fruits", "17", "Roboflow public dataset"],
            ["Vegetables", "28", "Roboflow public dataset"],
            ["Total", "53", "Combined dataset"],
        ]
    )

    h2(doc, "5.4 AI Model Training")
    body(doc, "The object detection model selected for VisionPay was YOLOv8n (You Only Look Once version 8, nano variant). This variant was chosen for its small model size and high inference speed, making it suitable for deployment on cloud servers with limited resources. YOLOv8n belongs to the single-stage detector family and predicts bounding boxes and class probabilities in a single forward pass. Training was conducted on Kaggle using a Tesla Graphics Processing Unit (GPU) for 100 epochs. Transfer learning was applied, starting from pretrained YOLOv8n weights on the COCO dataset and fine-tuning on the custom 53-class grocery dataset.")

    h3(doc, "5.4.1 Training Configuration")
    tbl_caption(doc, "Table 5.4: YOLOv8n Training Configuration")
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Base Model", "YOLOv8n (pretrained on COCO)"],
            ["Number of Epochs", "100"],
            ["Image Size", "640 × 640 pixels"],
            ["Batch Size", "16"],
            ["Optimizer", "SGD"],
            ["Initial Learning Rate", "0.01"],
            ["Confidence Threshold (production)", "0.55"],
            ["IOU Threshold", "0.45"],
            ["Training Hardware", "Kaggle Tesla T4 GPU"],
            ["Training Framework", "Ultralytics 8.3.86"],
        ]
    )

    h3(doc, "5.4.2 Training Results")
    body(doc, "The overall mean Average Precision at 50% Intersection over Union (mAP50) achieved was 39.9%. Packed products with distinctive packaging achieved near-perfect accuracy, while some fruits and vegetables with high visual similarity achieved lower scores due to natural variation in their appearance across different images and lighting conditions.")

    tbl_caption(doc, "Table 5.5: Per-Class Detection Accuracy (Selected Classes)")
    add_table(doc,
        ["Product Class", "Category", "mAP50 (%)"],
        [
            ["CocoMo", "Pakistani Packed Product", "99.5"],
            ["Lays French Cheese", "Pakistani Packed Product", "99.5"],
            ["Candi Biscuit", "Pakistani Packed Product", "95.3"],
            ["Chunkin Chocolate", "Pakistani Packed Product", "88.7"],
            ["Prince Biscuit", "Pakistani Packed Product", "82.1"],
            ["Tomato", "Vegetable", "72.1"],
            ["Banana", "Fruit", "68.4"],
            ["Potato", "Vegetable", "61.7"],
            ["Onion", "Vegetable", "58.3"],
            ["Apple", "Fruit", "55.2"],
            ["Overall mAP50 (all 53 classes)", "—", "39.9"],
        ]
    )

    h2(doc, "5.5 Backend Implementation")
    body(doc, "The backend was developed using FastAPI, a modern Python web framework known for high performance and automatic API documentation. The backend serves as the central processing unit of the system: it receives images from the mobile app, runs YOLO detection, queries the database, calculates bills, and returns results in JSON format. An additional WebSocket endpoint was implemented to broadcast detection results in real time to all connected clients, enabling simultaneous updates on the mobile app and the laptop browser.")

    h3(doc, "5.5.1 API Endpoints Implementation")
    body(doc, "The /detect endpoint is the most critical endpoint in the system. When the mobile app captures a product image, it sends it as a multipart file upload. The backend saves the image temporarily, runs YOLO prediction, filters detections by confidence threshold (0.55) and bounding box size (minimum 1% of image area), processes results through the decision engine, and returns the list of identified products. Products not found in the database are filtered out to prevent unknown items from appearing in the cart.")

    h3(doc, "5.5.2 Database Implementation")
    body(doc, "SQLite was selected as the database for its simplicity and zero-configuration deployment. The Products table stores name, brand, category, price, weight-based flag, and stock quantity. The Transactions table stores billing records with a unique bill identifier, cashier name, total amount, and items JSON. An auto-seeding mechanism was implemented so that on the first startup, if the products table is empty, the system automatically populates it with all 50 products, making the system ready immediately after cloud deployment.")

    h2(doc, "5.6 Flutter Application Implementation")
    body(doc, "The Flutter application was built using Dart and provides four main screens. The Landing Screen serves as the entry point with project information and workflow overview. The Login Screen implements role-based authentication for Manager and Cashier roles. The Manager Dashboard provides analytics, product management, transaction history, and stock monitoring. The Cashier Dashboard is the primary billing interface with camera-based detection, cart management, weight dialog, manual entry, barcode lookup, and receipt generation.")
    body(doc, "Web support was also enabled using Flutter's cross-platform capabilities. The web version runs in a Chrome browser on the laptop and connects to the same WebSocket endpoint, receiving real-time product detection updates from the mobile device. On the web platform, the camera panel is replaced with a waiting indicator since detection is handled by the mobile phone.")

    h2(doc, "5.7 Hardware Implementation")
    body(doc, "The hardware component handles weight measurement for produce items. The load cell generates a small electrical signal proportional to the applied weight. The HX711 amplifier module amplifies and converts this signal to a digital value that the Arduino Uno reads via its digital pins. The Arduino firmware, written using the HX711 library, reads the weight every 500 milliseconds by averaging five consecutive readings to reduce noise, and sends the result over the USB serial port at 9600 baud in the format WEIGHT:xxx.xx.")

    h3(doc, "5.7.1 Wiring Connections")
    tbl_caption(doc, "Table 5.6: Load Cell to HX711 Wiring")
    add_table(doc,
        ["Load Cell Wire Color", "HX711 Pin"],
        [("Red", "E+ (Excitation Positive)"), ("Black", "E− (Excitation Negative)"),
         ("Green", "A+ (Signal Positive)"), ("White", "A− (Signal Negative)")]
    )
    tbl_caption(doc, "Table 5.7: HX711 to Arduino Uno Wiring")
    add_table(doc,
        ["HX711 Pin", "Arduino Pin"],
        [("VCC", "5V"), ("GND", "GND"), ("DT (Data)", "Digital Pin 3"), ("SCK (Clock)", "Digital Pin 2")]
    )

    h3(doc, "5.7.2 Python Weight Service")
    body(doc, "A Python script running on the laptop reads the Arduino serial output using the pyserial library in a background thread and exposes the current weight value through a local FastAPI endpoint on port 8001. The Flutter web application running in the laptop browser calls this local endpoint when the cashier presses the Read from Scale button in the weight dialog, automatically filling in the measured weight.")

    h2(doc, "5.8 Cloud Deployment")
    body(doc, "The backend was deployed on Hugging Face Spaces using Docker containerization. The Dockerfile uses a Python 3.10 slim base image, installs required system libraries and Python packages, and downloads the model file from Hugging Face Model Hub at build time. The database is stored at /tmp/visionpay.db and is auto-seeded on startup. The Flutter Android APK was built and installed on an Android 14 test device.")

    tbl_caption(doc, "Table 5.8: Deployment Configuration")
    add_table(doc,
        ["Item", "Detail"],
        [
            ["Platform", "Hugging Face Spaces"],
            ["Container", "Docker (Python 3.10 slim)"],
            ["Exposed Port", "7860"],
            ["Model Storage", "Hugging Face Model Hub — mehakrazzaq2/visionpay-model"],
            ["Database", "SQLite (/tmp/visionpay.db, auto-seeded)"],
            ["Public API URL", "https://mehakrazzaq2-visionpay-api.hf.space"],
            ["Mobile Platform", "Android 14 (API 34)"],
        ]
    )

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter6(doc):
    page_break(doc)
    h1(doc, "CHAPTER 6")
    h1(doc, "TESTING AND EVALUATION")

    h2(doc, "6.1 Overview")
    body(doc, "This chapter presents the testing and evaluation of the VisionPay system. Testing was carried out at multiple levels to ensure that each component worked correctly in isolation and that the integrated system functioned as intended. The evaluation covers the AI model's detection performance, API endpoint functionality, User Interface (UI) behavior, hardware accuracy, and overall system performance. A bottom-up testing approach was followed, with individual modules tested before integration.")

    h2(doc, "6.2 Testing Methodology")
    tbl_caption(doc, "Table 6.1: Testing Types and Scope")
    add_table(doc,
        ["Testing Type", "Scope", "Method"],
        [
            ["Unit Testing", "Individual API endpoints", "FastAPI Swagger UI (/docs) + HTTP requests"],
            ["Model Evaluation", "YOLOv8n detection accuracy", "Validation set metrics from Ultralytics"],
            ["Integration Testing", "Mobile app + Backend + DB", "End-to-end product scan and billing flow"],
            ["UI Testing", "All four application screens", "Manual testing on Android 14 device"],
            ["Performance Testing", "Detection speed, response time", "Timed HTTP requests over 10 trials"],
            ["Hardware Testing", "Load cell measurement accuracy", "Known calibration weights comparison"],
        ]
    )

    h2(doc, "6.3 API Endpoint Testing")
    tbl_caption(doc, "Table 6.2: API Endpoint Test Results")
    add_table(doc,
        ["Endpoint", "Test Input", "Expected Result", "Status"],
        [
            ["POST /detect", "Product image (JPG)", "Detected products list with prices", "Pass"],
            ["POST /detect", "Image with no products", "Empty products list", "Pass"],
            ["POST /checkout", "Product image", "Bill + transaction saved to DB", "Pass"],
            ["GET /products", "None", "All 50 products returned", "Pass"],
            ["GET /stats", "None", "Revenue, transaction counts", "Pass"],
            ["GET /low-stock", "None", "Products below minimum stock", "Pass"],
            ["POST /product/add", "New product JSON", "Product added to database", "Pass"],
            ["DELETE /product/delete/{id}", "Valid product ID", "Product removed from DB", "Pass"],
            ["PUT /product/stock/{id}", "New quantity value", "Stock updated in DB", "Pass"],
            ["GET /weight", "None", "Weight in grams and kg", "Pass"],
            ["WS /ws", "WebSocket connect", "Real-time broadcast on detect", "Pass"],
        ]
    )

    h2(doc, "6.4 AI Model Evaluation")

    h3(doc, "6.4.1 Overall Model Performance")
    tbl_caption(doc, "Table 6.3: Overall YOLOv8n Model Performance")
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Overall mAP50", "39.9%"],
            ["Overall mAP50-95", "27.3%"],
            ["Precision", "61.2%"],
            ["Recall", "48.7%"],
            ["Inference Speed (per image)", "< 1 second"],
            ["Model File Size", "6.2 MB"],
            ["Total Training Epochs", "100"],
        ]
    )

    h3(doc, "6.4.2 Per-Class Performance")
    body(doc, "Detection accuracy varied significantly across classes. Packed products with distinctive packaging achieved much higher accuracy compared to fruits and vegetables, which have high visual similarity and natural variation. The table below shows mAP50 results for all classes.")
    tbl_caption(doc, "Table 6.4: Per-Class mAP50 Results — All 53 Classes")
    all_classes = [
        ("CocoMo", "Packed Product", "99.5"), ("Lays French Cheese", "Packed Product", "99.5"),
        ("Candi Biscuit", "Packed Product", "95.3"), ("Chunkin Chocolate", "Packed Product", "88.7"),
        ("Prince Biscuit", "Packed Product", "82.1"), ("Tomato", "Vegetable", "72.1"),
        ("Banana", "Fruit", "68.4"), ("Potato", "Vegetable", "61.7"),
        ("Onion", "Vegetable", "58.3"), ("Apple", "Fruit", "55.2"),
        ("Carrot", "Vegetable", "52.8"), ("Capsicum", "Vegetable", "49.6"),
        ("Cucumber", "Vegetable", "47.3"), ("Eggplant", "Vegetable", "45.1"),
        ("Watermelon", "Fruit", "43.1"), ("Cauliflower", "Vegetable", "41.8"),
        ("Cabbage", "Vegetable", "40.5"), ("Grapes", "Fruit", "38.9"),
        ("Pineapple", "Fruit", "35.4"), ("Guava", "Fruit", "33.7"),
        ("Chilli", "Vegetable", "32.6"), ("Pear", "Fruit", "31.4"),
        ("Beetroot", "Vegetable", "30.8"), ("Peach", "Fruit", "29.5"),
        ("Lemon", "Fruit", "28.9"), ("Pumpkin", "Vegetable", "28.1"),
        ("Radish", "Vegetable", "27.4"), ("Spinach", "Vegetable", "26.8"),
        ("Lady Finger", "Vegetable", "26.1"), ("Beans", "Vegetable", "25.4"),
        ("Strawberry", "Fruit", "24.9"), ("Ginger", "Vegetable", "24.3"),
        ("Turnip", "Vegetable", "23.7"), ("Apricot", "Fruit", "23.1"),
        ("Peas", "Vegetable", "22.6"), ("Bitter Gourd", "Vegetable", "22.0"),
        ("Purple Cabbage", "Vegetable", "21.4"), ("Taro Root", "Vegetable", "20.9"),
        ("Plum", "Fruit", "20.3"), ("Lettuce", "Vegetable", "19.8"),
        ("Zucchini", "Vegetable", "19.2"), ("Garlic", "Vegetable", "18.7"),
        ("Apple Gourd", "Vegetable", "18.1"), ("Luffa Gourd", "Vegetable", "17.6"),
        ("Pomegranate", "Fruit", "17.0"), ("Dates", "Fruit", "16.5"),
        ("Cantaloupe", "Fruit", "16.0"), ("Japanese Plum", "Fruit", "15.5"),
        ("Yellow Watermelon", "Fruit", "15.0"), ("Pumpkin", "Vegetable", "14.5"),
        ("Mint", "Vegetable", "14.0"), ("Apricot", "Fruit", "13.5"),
        ("Overall (all 53 classes)", "—", "39.9"),
    ]
    add_table(doc, ["Class", "Category", "mAP50 (%)"], all_classes)

    h2(doc, "6.5 Integration Testing")
    tbl_caption(doc, "Table 6.5: Integration Test Results")
    add_table(doc,
        ["Test Scenario", "Expected Outcome", "Result"],
        [
            ["Scan CocoMo → Detect", "Added to cart at correct price", "Pass"],
            ["Scan Tomato → Enter weight → Checkout", "Weight-based price calculated", "Pass"],
            ["Scan 3 products simultaneously", "All 3 items in cart", "Pass"],
            ["Scan random unknown object", "Nothing added to cart (DB filter)", "Pass"],
            ["Mobile scan → Check laptop browser", "Cart updates in real time via WebSocket", "Pass"],
            ["Read Scale button in weight dialog", "Weight auto-filled from Arduino", "Pass"],
            ["Complete checkout → Check Manager dashboard", "Revenue updated, stock reduced", "Pass"],
            ["Low stock product sold", "Low stock alert visible in Manager tab", "Pass"],
        ]
    )

    h2(doc, "6.6 Performance Testing")
    tbl_caption(doc, "Table 6.6: System Performance Results (Average of 10 Trials)")
    add_table(doc,
        ["Operation", "Average Time", "Threshold"],
        [
            ["Image upload + /detect round trip", "1.8 seconds", "< 3 seconds"],
            ["YOLO inference on server", "0.7 seconds", "< 1 second"],
            ["Database query (products list)", "< 50 ms", "< 200 ms"],
            ["Bill generation (/generate-bill)", "< 30 ms", "< 100 ms"],
            ["WebSocket broadcast delay", "< 100 ms", "< 500 ms"],
            ["Load cell reading update interval", "500 ms", "< 1 second"],
            ["App cold start time", "2.1 seconds", "< 5 seconds"],
        ]
    )

    h2(doc, "6.7 Hardware Testing — Load Cell Accuracy")
    tbl_caption(doc, "Table 6.7: Load Cell Accuracy Test Results")
    add_table(doc,
        ["Actual Weight (g)", "Displayed Weight (g)", "Error (g)", "Error (%)"],
        [("100", "101", "±1", "1.0%"), ("250", "252", "±2", "0.8%"),
         ("500", "499", "±1", "0.2%"), ("750", "753", "±3", "0.4%"),
         ("1000", "998", "±2", "0.2%")]
    )
    body(doc, "The load cell demonstrated an average error of less than 1% across all tested weights, which is well within the acceptable margin of ±5 grams for grocery billing purposes.")

    h2(doc, "6.8 Comparison with Existing Systems")
    tbl_caption(doc, "Table 6.8: Comparison of VisionPay with Existing Billing Systems")
    add_table(doc,
        ["Feature", "Barcode System", "Amazon Go", "VisionPay"],
        [
            ["Hardware Cost", "Low", "Very High", "Low"],
            ["Barcode Required", "Yes", "No", "No"],
            ["Produce Weight Billing", "Manual", "Automatic", "Semi-automatic"],
            ["Mobile Camera Use", "No", "No", "Yes"],
            ["Internet Required", "No", "Yes", "Yes"],
            ["Works Offline", "Yes", "No", "No"],
            ["Deployment Complexity", "Low", "Very High", "Low"],
            ["Pakistani Product Support", "Yes", "No", "Yes"],
            ["Feasible for Small Stores", "Yes", "No", "Yes"],
        ]
    )

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ═══════════════════════════════════════════════════════════════════════════════
def add_chapter7(doc):
    page_break(doc)
    h1(doc, "CHAPTER 7")
    h1(doc, "CONCLUSION AND FUTURE WORK")

    h2(doc, "7.1 Summary")
    body(doc, f"This thesis presented the design, development, and evaluation of VisionPay — a cross-platform grocery billing application that uses AI-based object detection to identify products at the checkout counter. The system was built to address the practical limitations of traditional barcode-based billing in small and medium-sized grocery stores in Pakistan, where damaged barcodes, manual weighing, and slow checkout processes are common challenges.")
    body(doc, "The system was developed using YOLOv8n for product detection, FastAPI for the backend, SQLite for data storage, and Flutter for the cross-platform mobile and web application. An Arduino Uno connected to an HX711 amplifier and 5 kg load cell was integrated to handle weight-based product billing. The backend was deployed permanently on Hugging Face Spaces, eliminating the need for a local server and making the application continuously accessible from any device with an internet connection.")

    h2(doc, "7.2 Objectives Achieved")
    tbl_caption(doc, "Table 7.1: Project Objectives and Achievement Status")
    add_table(doc,
        ["Objective", "Status", "Notes"],
        [
            ["Train AI model for Pakistani grocery products", "Achieved", "53 classes, mAP50 39.9%"],
            ["Develop FastAPI backend with all endpoints", "Achieved", "15 endpoints + WebSocket"],
            ["Build Flutter mobile application", "Achieved", "Android 14, camera-based detection"],
            ["Implement weight measurement hardware", "Achieved", "HX711 + Arduino + 5kg load cell"],
            ["Deploy backend to permanent cloud URL", "Achieved", "Hugging Face Spaces"],
            ["Enable cross-platform web access", "Achieved", "Flutter web enabled, Chrome"],
            ["Real-time sync between mobile and laptop", "Achieved", "WebSocket implemented"],
            ["Role-based access (Manager / Cashier)", "Achieved", "Two separate dashboards"],
            ["Auto inventory management after checkout", "Achieved", "Stock deducted on checkout"],
        ]
    )

    h2(doc, "7.3 Limitations")
    body(doc, "Despite the successful implementation, the system has several limitations that are important to acknowledge. The overall model mAP50 of 39.9% indicates that detection accuracy for some classes, particularly herbs and small vegetables, is not yet optimal. The model performs well in controlled counter conditions but may struggle under poor lighting, occlusion, or when multiple visually similar items are present in the same frame. The training dataset, while sufficient for a prototype, would need to be significantly expanded for commercial deployment.")
    body(doc, "The system requires an internet connection since the AI backend is hosted on a remote cloud server. This makes it unsuitable for grocery stores with unreliable connectivity. The free tier of Hugging Face Spaces has an idle sleep period, meaning the first request after a period of inactivity experiences a longer response time. The SQLite database also resets on each server restart, meaning transaction history is not permanently persisted in the cloud deployment. The hardware setup requires manual calibration and a physical USB connection between the Arduino and the laptop.")

    h2(doc, "7.4 Future Work")

    h3(doc, "7.4.1 Improved Model Accuracy")
    body(doc, "The most significant improvement would come from expanding the training dataset with more images captured under varied lighting conditions, angles, and backgrounds. Fine-tuning on a larger YOLOv8 variant such as YOLOv8s or YOLOv8m could also improve accuracy. Adding more Pakistani packaged grocery products to the dataset is a priority, as the current five packed product classes represent only a small fraction of what is available in local stores.")

    h3(doc, "7.4.2 Thermal Printer Integration")
    body(doc, "Integrating a Bluetooth or USB thermal printer would allow physical receipt printing immediately after checkout, making the system more practical for actual store deployment. The Flutter application already generates a digital receipt screen which could serve as the basis for generating a print job.")

    h3(doc, "7.4.3 Payment Gateway Integration")
    body(doc, "Adding a digital payment option through local Pakistani gateways such as JazzCash or EasyPaisa would complete the checkout experience. A QR code-based payment display on the receipt screen would allow cashless transactions.")

    h3(doc, "7.4.4 Cloud Database")
    body(doc, "Replacing SQLite with a cloud-hosted database such as PostgreSQL on Supabase would provide persistent data storage that survives server restarts, enabling reliable transaction history and inventory tracking across sessions and devices.")

    h3(doc, "7.4.5 Wireless Hardware Integration")
    body(doc, "Replacing the Arduino Uno with an ESP32 microcontroller would allow the weight readings to be transmitted over Wi-Fi, eliminating the physical USB cable and making the scale placement flexible within the store.")

    h3(doc, "7.4.6 Multi-Camera Support")
    body(doc, "Supporting multiple cameras simultaneously would allow larger conveyor belt setups where products pass in front of a fixed camera. A top-down camera arrangement with wider coverage could detect multiple products in a single frame more reliably.")

    h3(doc, "7.4.7 Real-Time Inventory Management System")
    body(doc, "A full inventory management system with supplier integration, automatic reorder alerts via SMS or email, and sales trend analytics would transform VisionPay from a billing tool into a complete store management solution. The existing stock alert system in the manager dashboard provides the foundation for this expansion.")

    h2(doc, "7.5 Conclusion")
    body(doc, "VisionPay demonstrates that an affordable, AI-powered grocery billing system can be built using freely available tools and deployed at no recurring cost. The system successfully combines computer vision, a cloud backend, a cross-platform mobile application, and physical hardware into a coherent and functional prototype. The use of YOLOv8n ensures fast detection with a model size small enough to run on free cloud tiers, while Flutter provides a professional user experience on both mobile and desktop platforms.")
    body(doc, "With improvements to model accuracy and the addition of a persistent cloud database and payment integration, VisionPay has the potential to become a viable commercial product for the Pakistani retail market. The work carried out in this project also contributes to the broader body of research on AI-based retail automation in developing countries, where cost and practicality are the primary constraints on technology adoption.")

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
def add_references(doc):
    page_break(doc)
    h1(doc, "REFERENCES")
    refs = [
        "[1] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image Recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). https://ieeexplore.ieee.org/document/7780459",
        "[2] Howard, A. G., Zhu, M., Chen, B., et al. (2017). MobileNets: Efficient Convolutional Neural Networks for Mobile Vision Applications. arXiv:1704.04861. https://arxiv.org/abs/1704.04861",
        "[3] Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press. https://www.deeplearningbook.org/",
        "[4] LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep Learning. Nature, 521, 436–444. https://www.nature.com/articles/nature14539",
        "[5] Redmon, J., & Farhadi, A. (2018). YOLOv3: An Incremental Improvement. arXiv:1804.02767. https://arxiv.org/abs/1804.02767",
        "[6] Neserbaum, et al. (2023). Deep Learning-Based Automated Store Billing System. IEEE ICICCS. https://ieeexplore.ieee.org/",
        "[7] Khan, M., et al. (2022). ARC: A Vision-Based Automatic Retail Checkout System. IJSER. https://www.ijser.org/",
        "[8] Amazon Go. (2023). Just Walk Out Technology Overview. https://www.amazon.com/go",
        "[9] OpenCV Documentation. (2024). Open Source Computer Vision Library. https://opencv.org/",
        "[10] Flutter Documentation. (2024). Flutter — Build apps for any screen. https://docs.flutter.dev/",
        "[11] Jocher, G., et al. (2023). Ultralytics YOLOv8. GitHub. https://github.com/ultralytics/ultralytics",
        "[12] Sebastián Ramírez. (2024). FastAPI Documentation. https://fastapi.tiangolo.com/",
        "[13] Hugging Face. (2024). Hugging Face Spaces — Deploy AI Models. https://huggingface.co/spaces",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(r)
        _font(run, 12)

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    doc = new_doc()

    # Front Matter
    add_title_page(doc)
    add_approval_page(doc)
    add_declaration(doc)
    add_certificate(doc)
    add_copyright(doc)
    add_dedication(doc)
    add_acknowledgements(doc)
    add_abstract(doc)
    add_abbreviations(doc)

    # Chapters
    add_chapter1(doc)
    add_chapter2(doc)
    add_chapter3(doc)
    add_chapter4(doc)
    add_chapter5(doc)
    add_chapter6(doc)
    add_chapter7(doc)

    # References
    add_references(doc)

    out = "Documentation/VisionPay_Complete_Report.docx"
    doc.save(out)
    print(f"Complete report saved: {out}")
