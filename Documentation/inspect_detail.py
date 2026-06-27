"""Detailed inspection around chapter 5 and 6 boundaries"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document(r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx")
paras = doc.paragraphs

def para_info(i, p):
    text = p.text.strip()[:80]
    style = p.style.name
    align = str(p.alignment)
    bold = any(r.bold for r in p.runs if r.text.strip())
    size = None
    for r in p.runs:
        if r.font.size:
            size = r.font.size.pt
            break
    # Check for page break
    has_pb = '<w:lastRenderedPageBreak' in p._element.xml or \
             'w:type="page"' in p._element.xml or \
             'w:pageBreakBefore' in p._element.xml
    has_pb_run = any('w:br' in r._element.xml and 'page' in r._element.xml for r in p.runs)
    return f"[{i:3}] sty='{style}' align={align} bold={bold} size={size} pb={has_pb or has_pb_run} | '{text}'"

# Show paragraphs around chapter 5 title (670-695)
print("=== CHAPTER 5 AREA (670-700) ===")
for i, p in enumerate(paras):
    if 670 <= i <= 700:
        print(para_info(i, p))

print("\n=== CHAPTER 6 AREA (755-785) ===")
for i, p in enumerate(paras):
    if 755 <= i <= 785:
        print(para_info(i, p))

# Also check tables count in doc
print(f"\nTotal tables in document: {len(doc.tables)}")
# Show tables near chapter 5 and 6
# We need to check the doc body elements
from docx.oxml.ns import nsmap
body = doc.element.body
elements = list(body)
print(f"Total body elements: {len(elements)}")
# Find table elements and their position
tbl_positions = []
for idx, el in enumerate(elements):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag == 'tbl':
        tbl_positions.append(idx)
print(f"Table positions in body: {tbl_positions[:30]}...")
