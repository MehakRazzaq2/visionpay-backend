"""Verify final state: all captions and LOT entries"""
from docx import Document
import re

doc = Document(r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx")
paras = doc.paragraphs

print("=== All Caption paragraphs (Tables only) ===")
chapter = 0
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^(?:CHAPTER|Chapter)\s+(\d)', t)
    if m:
        chapter = int(m.group(1))
    if p.style.name == 'Caption' and t.lower().startswith('table'):
        print(f"  Ch{chapter} [{i:4}] {t}")

print("\n=== List of Tables entries ===")
in_lot = False
for i, p in enumerate(paras):
    t = p.text.strip()
    if 'List Of Tables' in t or 'LIST OF TABLES' in t:
        in_lot = True
        print(f"  LOT HEADING: '{t}'")
        continue
    if in_lot and p.style.name == 'table of figures' and t.startswith('Table'):
        print(f"  {t}")
    elif in_lot and p.style.name not in ('table of figures', 'Normal', '') and t:
        break  # end of LOT

print(f"\nTotal paragraphs: {len(paras)}")
print(f"Total tables: {len(doc.tables)}")

# Quick check: new tables created?
print("\n=== Tables around Ch5.8 area ===")
for i, p in enumerate(paras):
    t = p.text.strip()
    if '5.8' in t or 'Deployment' in t or 'Cloud Deploy' in t:
        print(f"  [{i:4}] style={p.style.name} | {t[:80]}")

print("\n=== Tables around Ch6.2 area ===")
for i, p in enumerate(paras):
    t = p.text.strip()
    if '6.2' in t or 'API Endpoint Test' in t:
        print(f"  [{i:4}] style={p.style.name} | {t[:80]}")

print("\n=== Tables around Ch7.1 area ===")
for i, p in enumerate(paras):
    t = p.text.strip()
    if '7.1' in t or 'Objectives and Achievement' in t:
        print(f"  [{i:4}] style={p.style.name} | {t[:80]}")
