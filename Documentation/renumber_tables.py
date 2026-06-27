"""
Complete table renumbering and LOT rebuild for FYP Report - Copy.docx
- Assigns Table X.Y numbers by chapter
- Updates/adds captions before each table
- Rebuilds List of Tables with correct page numbers
"""
import re, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"

# ── Known table descriptions keyed by (chapter, sequential_number) ────────────
# These match the table content - identified from earlier doc.tables inspection
TABLE_DESCRIPTIONS = {
    # Front matter / system tables - skip (committee, abbreviations)
    # Chapter 2
    (2, 1): "Literature Review Summary",
    # Chapter 4 use cases
    (4, 1):  "Use Case — Place Product on Counter",
    (4, 2):  "Use Case — Capture Product Image",
    (4, 3):  "Use Case — Measure Product Weight",
    (4, 4):  "Use Case — Recognize Product",
    (4, 5):  "Use Case — Detect Brand",
    (4, 6):  "Use Case — Extract Text Using OCR",
    (4, 7):  "Use Case — Fetch Price from Database",
    (4, 8):  "Use Case — Validate Product Using Weight",
    (4, 9):  "Use Case — Calculate Total Bill",
    (4, 10): "Use Case — Generate Invoice",
    (4, 11): "Use Case — Make Payment",
    (4, 12): "Use Case — Update Inventory Automatically",
    (4, 13): "Use Case — Manage Store Operations",
    # Chapter 5
    (5, 1): "Development Tools and Libraries",
    (5, 2): "Hardware Components",
    (5, 3): "Dataset Summary",
    (5, 4): "YOLOv8n Training Configuration",
    (5, 5): "Per-Class Detection Accuracy (Selected Classes)",
    (5, 6): "Load Cell to HX711 Wiring",
    (5, 7): "HX711 to Arduino Wiring",
    (5, 8): "Deployment Configuration",
    # Chapter 6
    (6, 1): "Testing Types and Scope",
    (6, 2): "API Endpoint Test Results",
    (6, 3): "Overall Model Performance Metrics",
    (6, 4): "Per-Class mAP50 Results",
    (6, 5): "Integration Test Scenarios",
    (6, 6): "System Performance Results",
    (6, 7): "Load Cell Accuracy Test Results",
    (6, 8): "Comparison with Existing Billing Systems",
    # Chapter 7
    (7, 1): "Project Objectives and Achievement Status",
}

# Tables to SKIP (not chapter content: approval/committee table, abbreviations)
SKIP_TABLE_INDICES = {0, 1, 2}  # first 3 tables are front-matter tables

def get_text_from_el(el):
    """Get all text from an XML element"""
    parts = []
    for t in el.iter(qn('w:t')):
        if t.text:
            parts.append(t.text)
    return ''.join(parts).strip()

def make_caption_para(doc, caption_text):
    """Create a properly formatted Caption paragraph"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Style
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Caption')
    pPr.append(pStyle)
    # Centered
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    r.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = caption_text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_el)
    p.append(r)
    return p

# ── PASS 1: Walk body XML, track chapters, process tables ────────────────────
doc = Document(DOC_PATH)
body = doc.element.body
body_children = list(body)

chapter = 0
table_seq = {}       # chapter -> next sequential number
table_captions = []  # list of (chapter, num, description, table_el, prev_el)
table_global_idx = 0

print("=== Scanning document body for tables ===")

for idx, el in enumerate(body_children):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag

    if tag == 'p':
        text = get_text_from_el(el)
        # Detect chapter boundaries
        m = re.match(r'^(?:CHAPTER|Chapter)\s+(\d)', text)
        if m:
            chapter = int(m.group(1))

    elif tag == 'tbl':
        table_global_idx += 1

        # Skip front-matter tables
        if table_global_idx - 1 in SKIP_TABLE_INDICES:
            print(f"  Skipping table #{table_global_idx} (front matter, Ch{chapter})")
            continue

        # Skip tables in chapter 0 (front matter)
        if chapter == 0:
            print(f"  Skipping table #{table_global_idx} (Ch0 - front matter)")
            continue

        # Also skip Ch3 (no tables expected there)
        if chapter == 3:
            print(f"  Skipping table #{table_global_idx} (Ch3 - no chapter tables)")
            continue

        # Assign sequential number
        seq = table_seq.get(chapter, 0) + 1
        table_seq[chapter] = seq
        key = (chapter, seq)

        desc = TABLE_DESCRIPTIONS.get(key, f"Table {chapter}.{seq}")
        caption_text = f"Table {chapter}.{seq}: {desc}"

        # Check if previous element is already a Caption paragraph
        prev_el = body_children[idx - 1] if idx > 0 else None
        prev_is_caption = False
        if prev_el is not None and prev_el.tag.split('}')[-1] == 'p':
            prev_style_el = prev_el.find('.//' + qn('w:pStyle'))
            if prev_style_el is not None:
                prev_style = prev_style_el.get(qn('w:val'), '')
                if 'Caption' in prev_style or 'caption' in prev_style.lower():
                    prev_is_caption = True

        table_captions.append((chapter, seq, caption_text, el, prev_el, prev_is_caption))
        print(f"  Ch{chapter} Table #{seq}: '{caption_text}' | prev_is_caption={prev_is_caption}")

print(f"\nTotal chapter tables found: {len(table_captions)}")

# ── PASS 2: Update or add caption paragraphs ─────────────────────────────────
print("\n=== Updating/adding captions ===")

for (chapter, seq, caption_text, tbl_el, prev_el, prev_is_caption) in table_captions:
    if prev_is_caption and prev_el is not None:
        # Update existing caption text
        # Clear all runs in the caption paragraph and rewrite
        for r_el in list(prev_el.findall('.//' + qn('w:r'))):
            prev_el.remove(r_el)
        # Add new run with correct text
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rf)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.text = caption_text
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t_el)
        prev_el.append(r)
        print(f"  UPDATED caption: {caption_text}")
    else:
        # Insert NEW caption paragraph before the table
        new_cap = make_caption_para(doc, caption_text)
        tbl_el.addprevious(new_cap)
        print(f"  ADDED caption: {caption_text}")

# ── PASS 3: Fix remaining Caption paragraphs with old numbering ──────────────
# Find and fix any Caption paragraphs that still have old Table X format
# (these might be captions that WERE preceding tables but we already updated them above,
#  plus any stragglers)
print("\n=== Fixing stray old-format captions ===")
doc_check = Document(DOC_PATH)  # re-read isn't possible - work with in-memory doc
paras_check = doc.paragraphs
old_cap_pattern = re.compile(r'^Table\s+(?:\d{1,2}|[56]\.\d|7\.\d)(?:\s|:)')

# Also remove duplicate/empty Caption paragraphs
for p in paras_check:
    if p.style.name == 'Caption':
        t = p.text.strip()
        # Empty captions - clear them (they're blank figure placeholders, leave them)
        if not t:
            continue
        # Old "Table 2 Use Case 1" style - these should already be updated
        # But check if any got missed
        if old_cap_pattern.match(t) and not re.match(r'^Table \d+\.\d+:', t):
            print(f"  Still old format: '{t}' - skipping (should be handled)")

# ── PASS 4: Get page numbers via win32com and rebuild LOT ────────────────────
doc.save(DOC_PATH)
print(f"\nSaved intermediate: {DOC_PATH}")

# Now get page numbers
import win32com.client
wdActiveEndPageNumber = 3

word = win32com.client.Dispatch("Word.Application")
word.Visible = False

caption_page_map = {}  # "Table X.Y: Description" -> page number

try:
    doc_com = word.Documents.Open(DOC_PATH)
    doc_com.Repaginate()

    target_prefixes = [f"Table {ch}.{seq}:" for (ch, seq, _, __, ___, ____) in table_captions]

    for para in doc_com.Paragraphs:
        text = para.Range.Text.strip()
        for prefix in target_prefixes:
            if text.startswith(prefix):
                pg = para.Range.Information(wdActiveEndPageNumber)
                caption_page_map[text] = pg
                break

    doc_com.Close(False)
    print(f"\nPage numbers found: {len(caption_page_map)}")
    for k, v in caption_page_map.items():
        print(f"  {k} -> page {v}")
finally:
    word.Quit()

# ── PASS 5: Rebuild List of Tables ────────────────────────────────────────────
doc2 = Document(DOC_PATH)
paras2 = doc2.paragraphs

# Find the List of Tables heading and its last entry
lot_heading_idx = None
lot_entries_start = None
lot_entries_end = None

for i, p in enumerate(paras2):
    t = p.text.strip()
    if 'List Of Tables' in t and p.style.name == 'Normal' and lot_heading_idx is None:
        lot_heading_idx = i
    if lot_heading_idx is not None and p.style.name == 'table of figures' and 'Table' in p.text:
        if lot_entries_start is None:
            lot_entries_start = i
        lot_entries_end = i

print(f"\nLOT heading at [{lot_heading_idx}], entries [{lot_entries_start}]...[{lot_entries_end}]")

# Remove all existing LOT table entries
if lot_entries_start is not None:
    to_remove = []
    for i, p in enumerate(paras2):
        if lot_entries_start <= i <= lot_entries_end and p.style.name == 'table of figures' and 'Table' in p.text:
            to_remove.append(p._element)
    for el in to_remove:
        el.getparent().remove(el)
    print(f"Removed {len(to_remove)} old LOT entries")

# Re-read after removal
doc2.save(DOC_PATH)
doc3 = Document(DOC_PATH)
paras3 = doc3.paragraphs

# Find insertion point (after last LOT figure entry, before abbreviations)
lot_last_fig_el = None
for i, p in enumerate(paras3):
    if p.style.name == 'table of figures':
        lot_last_fig_el = p._element

# Build new LOT entries in order
def make_lot_entry_el(text, page_num):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'TableofFigures')
    pPr.append(pStyle)
    p.append(pPr)
    # Text run
    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    rf1 = OxmlElement('w:rFonts'); rf1.set(qn('w:ascii'), 'Times New Roman'); rf1.set(qn('w:hAnsi'), 'Times New Roman')
    rPr1.append(rf1)
    sz1 = OxmlElement('w:sz'); sz1.set(qn('w:val'), '24'); rPr1.append(sz1)
    r1.append(rPr1)
    t1 = OxmlElement('w:t'); t1.text = text; t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r1.append(t1); p.append(r1)
    # Tab
    r2 = OxmlElement('w:r'); r2.append(OxmlElement('w:tab')); p.append(r2)
    # Page number
    r3 = OxmlElement('w:r')
    rPr3 = OxmlElement('w:rPr')
    rf3 = OxmlElement('w:rFonts'); rf3.set(qn('w:ascii'), 'Times New Roman'); rf3.set(qn('w:hAnsi'), 'Times New Roman')
    rPr3.append(rf3)
    sz3 = OxmlElement('w:sz'); sz3.set(qn('w:val'), '24'); rPr3.append(sz3)
    r3.append(rPr3)
    t3 = OxmlElement('w:t'); t3.text = str(page_num); r3.append(t3); p.append(r3)
    return p

# Insert all table LOT entries after the last figure entry
insert_after = lot_last_fig_el
for (ch, seq, caption_text, _, __, ___) in table_captions:
    lot_text = f"Table {ch}.{seq}: {TABLE_DESCRIPTIONS.get((ch, seq), 'Table')}"
    pg = caption_page_map.get(caption_text, '-')
    new_entry = make_lot_entry_el(lot_text, pg)
    insert_after.addnext(new_entry)
    insert_after = new_entry
    print(f"  LOT entry: {lot_text} -> p{pg}")

doc3.save(DOC_PATH)
print(f"\nDONE. Saved: {DOC_PATH}")
