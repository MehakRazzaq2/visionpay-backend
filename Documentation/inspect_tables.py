"""Check actual table contents in Ch5 and Ch6"""
from docx import Document

doc = Document(r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx")

# Print ALL paragraphs (with and without text) between indices 685 and 760
# including table cells
paras = doc.paragraphs
print("=== ALL PARAS 685-770 (including empty) ===")
for i, p in enumerate(paras):
    if 685 <= i <= 770:
        style = p.style.name
        text = p.text[:80] if p.text else ''
        print(f"[{i:3}] '{style}' | '{text}'")

print("\n=== TABLES in doc ===")
for ti, tbl in enumerate(doc.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns) if tbl.rows else 0
    # Get first few cells
    sample = []
    for r in tbl.rows[:2]:
        for c in r.cells:
            sample.append(c.text[:20])
    print(f"Table[{ti}] {rows}rows x {cols}cols | first cells: {sample[:6]}")
