"""Deep inspection - find ALL paragraphs near expected table captions"""
from docx import Document
import re

doc = Document(r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx")
paras = doc.paragraphs

# Search for any paragraph mentioning Table 5.8 or Table 6.2 anywhere
print("=== Search for Table 5.8, 6.2, Literature Review caption ===")
for i, p in enumerate(paras):
    t = p.text.strip()
    if any(x in t for x in ["5.8", "6.2", "Literature Review", "API Endpoint", "Deployment Config"]):
        print(f"[{i:3}] style='{p.style.name}' | '{t[:80]}'")

# Show all paragraphs in Ch5 area near tables 5.6-5.8
print("\n=== Ch5 area 780-810 (all paragraphs) ===")
for i, p in enumerate(paras):
    if 780 <= i <= 810:
        t = p.text.strip()
        print(f"[{i:3}] style='{p.style.name}' | '{t[:80]}'")

# Show Ch6 area 820-860
print("\n=== Ch6 area 820-875 (all paragraphs) ===")
for i, p in enumerate(paras):
    if 820 <= i <= 875:
        t = p.text.strip()
        print(f"[{i:3}] style='{p.style.name}' | '{t[:80]}'")

# Chapter 2 area for Literature Review table
print("\n=== Ch2 area 380-415 (all paragraphs) ===")
for i, p in enumerate(paras):
    if 380 <= i <= 415:
        t = p.text.strip()
        print(f"[{i:3}] style='{p.style.name}' | '{t[:80]}'")

print(f"\nTotal paragraphs in document: {len(paras)}")
print(f"Total tables in document: {len(doc.tables)}")
