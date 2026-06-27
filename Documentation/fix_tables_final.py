"""
Comprehensive table fix for FYP Report - Copy.docx:
Phase 1: Create 3 new Word tables from bullet-text sections (5.8, 6.2, 7.1)
Phase 2: Renumber all chapter table captions sequentially
Phase 3: Get page numbers via win32com
Phase 4: Rebuild List of Tables
"""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"C:\Users\Mehak Razzaq\Desktop\VisionPay\Documentation\FYP Report - Copy.docx"

# ── Table data for newly created tables ──────────────────────────────────────
T58_HEADERS = ["Component", "Details"]
T58_ROWS = [
    ["Backend",   "Docker container on Hugging Face Spaces, Python 3.10"],
    ["AI Model",  "YOLOv8n, stored on Hugging Face Model Hub, downloaded at build time"],
    ["Database",  "SQLite stored at /tmp/visionpay.db on the server"],
    ["API URL",   "https://mehakrazzaq2-visionpay-api.hf.space"],
    ["Mobile App","Flutter targeting Android 14 (API level 34)"],
]
T58_WIDTHS = [2500, 6860]

T62_HEADERS = ["Endpoint", "Method", "Test Result"]
T62_ROWS = [
    ["/detect",              "POST",   "Returned product name, confidence score, and price"],
    ["/checkout",            "POST",   "Generated bill with product details and total amount"],
    ["/products",            "GET",    "Returned all 50 seeded products from the database"],
    ["/stats",               "GET",    "Returned revenue figures and transaction counts"],
    ["/low-stock",           "GET",    "Returned list of products below stock threshold"],
    ["/product/add",         "POST",   "New product successfully added to the database"],
    ["/product/delete/{id}", "DELETE", "Product successfully removed by valid ID"],
    ["/weight",              "GET",    "Returned weight reading in grams and kilograms"],
    ["/ws",                  "WS",     "WebSocket established; real-time broadcasts received"],
]
T62_WIDTHS = [2200, 1200, 5960]

T71_HEADERS = ["Objective", "Status"]
T71_ROWS = [
    ["AI model trained for 53 Pakistani grocery product classes",            "Achieved"],
    ["FastAPI backend with 15 fully functional endpoints",                   "Achieved"],
    ["Flutter mobile app for Android 14 with camera-based billing",          "Achieved"],
    ["Hardware weight measurement using HX711, Arduino Uno, and load cell",  "Achieved"],
    ["Permanent cloud deployment on Hugging Face Spaces",                    "Achieved"],
    ["Cross-platform web access via Flutter web",                            "Achieved"],
    ["Real-time mobile-to-laptop sync via WebSocket",                        "Achieved"],
    ["Role-based access with Manager and Cashier dashboards",                "Achieved"],
    ["Automatic stock deduction after each billing transaction",             "Achieved"],
]
T71_WIDTHS = [7360, 2000]

# ── Chapter 4 use-case descriptions (in document order) ──────────────────────
CH4_DESCRIPTIONS = [
    "Use Case — Place Product on Counter",
    "Use Case — Capture Product Image",
    "Use Case — Measure Product Weight",
    "Use Case — Recognize Product",
    "Use Case — Detect Brand",
    "Use Case — Extract Text Using OCR",
    "Use Case — Fetch Price from Database",
    "Use Case — Validate Product Using Weight",
    "Use Case — Calculate Total Bill",
    "Use Case — Generate Invoice",
    "Use Case — Make Payment",
    "Use Case — Update Inventory Automatically",
    "Use Case — Manage Store Operations",
]

# ── XML helpers ───────────────────────────────────────────────────────────────
def xml_space(val='preserve'):
    return '{http://www.w3.org/XML/1998/namespace}space'

def make_run(text, bold=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
        b2 = OxmlElement('w:bCs'); rPr.append(b2)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(xml_space(), 'preserve')
    r.append(t)
    return r

def make_cell(text, bold=False, width=None, shading=None):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if width:
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
    if shading:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shading)
        tcPr.append(shd)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'left'); pPr.append(jc)
    p.append(pPr)
    p.append(make_run(text, bold=bold))
    tc.append(p)
    return tc

def make_table_xml(headers, rows, col_widths=None):
    """Build a properly formatted Word table XML element"""
    tbl = OxmlElement('w:tbl')

    # Table properties: bordered, full width
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle'); tblStyle.set(qn('w:val'), 'TableGrid'); tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Column grid
    tblGrid = OxmlElement('w:tblGrid')
    if col_widths:
        for w in col_widths:
            gridCol = OxmlElement('w:gridCol'); gridCol.set(qn('w:w'), str(w)); tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def make_row(cell_texts, bold_cells=False, widths=None, shading=None):
        tr = OxmlElement('w:tr')
        for i, ct in enumerate(cell_texts):
            w = widths[i] if widths else None
            tc = make_cell(ct, bold=bold_cells, width=w, shading=shading)
            tr.append(tc)
        return tr

    # Header row with light shading
    tr_hdr = make_row(headers, bold_cells=True, widths=col_widths, shading='D9D9D9')
    tbl.append(tr_hdr)

    # Data rows
    for row in rows:
        tr = make_row(row, bold_cells=False, widths=col_widths)
        tbl.append(tr)

    return tbl

def make_caption_el(caption_text):
    """Create a Caption-style paragraph element"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), 'Caption'); pPr.append(pStyle)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
    p.append(pPr)
    p.append(make_run(caption_text))
    return p

def get_text(el):
    parts = []
    for t in el.iter(qn('w:t')):
        if t.text:
            parts.append(t.text)
    return ''.join(parts).strip()

def get_para_style(el):
    s = el.find('.//' + qn('w:pStyle'))
    return s.get(qn('w:val'), 'Normal') if s is not None else 'Normal'

# ── LOAD DOCUMENT ─────────────────────────────────────────────────────────────
doc = Document(DOC_PATH)
body = doc.element.body

print("=== Phase 1: Creating missing tables from text content ===\n")

# ── 1a. Create Table 5.8 (Deployment Configuration) ──────────────────────────
# Find: "Key deployment details are as follows:" paragraph
# Then find the 5 bullet paragraphs that follow it
intro_58 = None
bullets_58 = []
found_key = False
for el in list(body):
    if el.tag.split('}')[-1] != 'p':
        continue
    txt = get_text(el)
    if txt == 'Key deployment details are as follows:':
        intro_58 = el
        found_key = True
        continue
    if found_key and len(bullets_58) < 5:
        if txt and el.tag.split('}')[-1] == 'p':
            bullets_58.append(el)
    elif found_key and len(bullets_58) == 5:
        break

if intro_58 is not None:
    cap_58 = make_caption_el("Table 5.8: Deployment Configuration")
    tbl_58 = make_table_xml(T58_HEADERS, T58_ROWS, T58_WIDTHS)
    # Insert caption and table AFTER the intro paragraph
    intro_58.addnext(tbl_58)
    intro_58.addnext(cap_58)
    # Remove original bullet paragraphs
    for b in bullets_58:
        b.getparent().remove(b)
    print("  Created Table 5.8: Deployment Configuration (5 rows)")
else:
    print("  WARNING: Could not find 'Key deployment details' paragraph for Table 5.8")

# ── 1b. Create Table 6.2 (API Endpoint Test Results) ─────────────────────────
# Find the intro paragraph of section 6.3 API Endpoint Testing
# then find the 9 bullet point paragraphs

# Look for body heading "6.3." and then the intro Normal paragraph
# The intro is: "Every endpoint was tested through the Swagger UI at /docs..."
# The bullets start: "POST /detect: was tested..."
intro_62 = None
bullets_62 = []
for el in list(body):
    if el.tag.split('}')[-1] != 'p':
        continue
    txt = get_text(el)
    if txt.startswith('Every endpoint was tested through the Swagger UI'):
        intro_62 = el
        continue
    if intro_62 is not None and len(bullets_62) < 9:
        if txt and txt.startswith(('POST', 'GET', 'DELETE', 'WS')):
            bullets_62.append(el)
    elif intro_62 is not None and len(bullets_62) == 9:
        break

if intro_62 is not None:
    cap_62 = make_caption_el("Table 6.2: API Endpoint Test Results")
    tbl_62 = make_table_xml(T62_HEADERS, T62_ROWS, T62_WIDTHS)
    intro_62.addnext(tbl_62)
    intro_62.addnext(cap_62)
    for b in bullets_62:
        b.getparent().remove(b)
    print(f"  Created Table 6.2: API Endpoint Test Results ({len(T62_ROWS)} rows)")
else:
    print("  WARNING: Could not find API endpoint section paragraphs for Table 6.2")

# ── 1c. Create Table 7.1 (Project Objectives and Achievement Status) ─────────
# Find: "The following objectives set at the beginning..."
# Then find 9 objectives list paragraphs
intro_71 = None
bullets_71 = []
for el in list(body):
    if el.tag.split('}')[-1] != 'p':
        continue
    txt = get_text(el)
    if txt.startswith('The following objectives set at the beginning of the project were'):
        intro_71 = el
        continue
    if intro_71 is not None and len(bullets_71) < 9:
        if txt and len(txt) > 20:
            style = get_para_style(el)
            if style in ('Normal', 'ListParagraph') or 'Normal' in style:
                bullets_71.append(el)
    elif intro_71 is not None and len(bullets_71) == 9:
        break

if intro_71 is not None and bullets_71:
    cap_71 = make_caption_el("Table 7.1: Project Objectives and Achievement Status")
    tbl_71 = make_table_xml(T71_HEADERS, T71_ROWS, T71_WIDTHS)
    intro_71.addnext(tbl_71)
    intro_71.addnext(cap_71)
    for b in bullets_71:
        b.getparent().remove(b)
    print(f"  Created Table 7.1: Project Objectives and Achievement Status ({len(T71_ROWS)} rows)")
else:
    print(f"  WARNING: Could not find objectives paragraphs for Table 7.1 (intro={intro_71 is not None}, bullets={len(bullets_71)})")

# ── PHASE 2: Renumber all chapter table captions ──────────────────────────────
print("\n=== Phase 2: Renumbering all chapter table captions ===\n")

chapter = 0
ch_table_seq = {}  # chapter -> sequential number for this pass
ch4_desc_idx = 0   # index into CH4_DESCRIPTIONS

# Chapter 5 descriptions (in document order)
CH5_DESCRIPTIONS = [
    "Development Tools and Libraries",
    "Hardware Components Used",
    "Dataset Summary",
    "YOLOv8n Training Configuration",
    "Per-Class Detection Accuracy (Selected Classes)",
    "Load Cell to HX711 Wiring",
    "HX711 to Arduino Wiring",
    "Deployment Configuration",
]

# Chapter 6 descriptions - sequential for tables that EXIST
CH6_DESCRIPTIONS = [
    "Testing Types and Scope",
    "API Endpoint Test Results",
    "Overall Model Performance Metrics",
    "Integration Test Scenarios",
    "System Performance Results",
    "Load Cell Accuracy Test Results",
    "Comparison with Existing Billing Systems",
]

# Chapter 7 descriptions
CH7_DESCRIPTIONS = [
    "Project Objectives and Achievement Status",
]

CHAPTER_DESCS = {
    4: CH4_DESCRIPTIONS,
    5: CH5_DESCRIPTIONS,
    6: CH6_DESCRIPTIONS,
    7: CH7_DESCRIPTIONS,
}

SKIP_CHAPTER = {0, 1, 2, 3}  # no chapter tables in these
front_matter_tbl_count = 0
MAX_FRONT_MATTER_TABLES = 3  # skip first 3 tables (committee, etc.)

updated_captions = []  # (caption_text, tbl_el) for LOT

for el in list(body):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag

    if tag == 'p':
        txt = get_text(el)
        m = re.match(r'^(?:CHAPTER|Chapter)\s+(\d)', txt)
        if m:
            chapter = int(m.group(1))

    elif tag == 'tbl':
        # Skip front-matter tables
        if chapter in SKIP_CHAPTER or front_matter_tbl_count < MAX_FRONT_MATTER_TABLES:
            front_matter_tbl_count += 1
            continue

        seq = ch_table_seq.get(chapter, 0) + 1
        ch_table_seq[chapter] = seq

        descs = CHAPTER_DESCS.get(chapter, [])
        desc = descs[seq - 1] if seq - 1 < len(descs) else f"Table {chapter}.{seq}"
        new_caption_text = f"Table {chapter}.{seq}: {desc}"

        # Find the preceding element (should be Caption or we need to add one)
        prev = el.getprevious()
        if prev is not None:
            prev_tag = prev.tag.split('}')[-1] if '}' in prev.tag else ''
            if prev_tag == 'p':
                prev_style = get_para_style(prev)
                if 'Caption' in prev_style or 'caption' in prev_style.lower():
                    # Update existing caption
                    # Remove all w:r elements and rebuild
                    for r_el in list(prev.findall('.//' + qn('w:r'))):
                        r_el.getparent().remove(r_el)
                    prev.append(make_run(new_caption_text))
                    # Ensure centered
                    pPr = prev.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr'); prev.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc'); pPr.append(jc)
                    jc.set(qn('w:val'), 'center')
                    print(f"  UPDATED: {new_caption_text}")
                    updated_captions.append((new_caption_text, el))
                else:
                    # Previous element is not a caption - add one
                    new_cap = make_caption_el(new_caption_text)
                    el.addprevious(new_cap)
                    print(f"  ADDED:   {new_caption_text}")
                    updated_captions.append((new_caption_text, el))
            else:
                new_cap = make_caption_el(new_caption_text)
                el.addprevious(new_cap)
                print(f"  ADDED:   {new_caption_text}")
                updated_captions.append((new_caption_text, el))
        else:
            new_cap = make_caption_el(new_caption_text)
            el.addprevious(new_cap)
            print(f"  ADDED:   {new_caption_text}")
            updated_captions.append((new_caption_text, el))

print(f"\nTotal captions updated/added: {len(updated_captions)}")

# ── Save after Phase 1+2 ─────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\nSaved intermediate (Phases 1+2): {DOC_PATH}")

# ── PHASE 3: Get page numbers via win32com ────────────────────────────────────
print("\n=== Phase 3: Getting page numbers via Word COM ===\n")
import win32com.client
wdActiveEndPageNumber = 3

word = win32com.client.Dispatch("Word.Application")
word.Visible = False
caption_page_map = {}

try:
    doc_com = word.Documents.Open(DOC_PATH)
    doc_com.Repaginate()

    target_caps = [ct for (ct, _) in updated_captions]

    for para in doc_com.Paragraphs:
        text = para.Range.Text.strip()
        for target in target_caps:
            if text == target:
                pg = para.Range.Information(wdActiveEndPageNumber)
                caption_page_map[target] = pg
                break

    doc_com.Close(False)
    print(f"Page numbers retrieved: {len(caption_page_map)}")
    for k, v in sorted(caption_page_map.items()):
        print(f"  {k} -> p{v}")
finally:
    word.Quit()

# ── PHASE 4: Rebuild List of Tables ──────────────────────────────────────────
print("\n=== Phase 4: Rebuilding List of Tables ===\n")

doc2 = Document(DOC_PATH)
paras2 = doc2.paragraphs

# Find the "List Of Tables" heading
lot_heading_idx = None
for i, p in enumerate(paras2):
    t = p.text.strip()
    if ('List Of Tables' in t or 'LIST OF TABLES' in t) and lot_heading_idx is None:
        lot_heading_idx = i
        print(f"  LOT heading found at para[{i}]: '{t}'")
        break

if lot_heading_idx is None:
    print("  ERROR: Could not find 'List Of Tables' heading!")
else:
    # Find all existing 'table of figures' style entries that are for Tables (not Figures)
    to_remove = []
    last_lot_el = None
    in_lot = False
    for i, p in enumerate(paras2):
        if i == lot_heading_idx:
            in_lot = True
            continue
        if in_lot:
            if p.style.name == 'table of figures':
                t = p.text.strip()
                if t.startswith('Table'):
                    to_remove.append(p._element)
                    print(f"  Removing old entry: '{t[:60]}'")
                elif t.startswith('Figure') or not t.startswith('Table'):
                    last_lot_el = p._element  # last figure entry
            # Stop when we hit something that's clearly not LOT content
            elif p.style.name not in ('table of figures', 'Normal') and p.style.name != '':
                if p.style.name.startswith('Heading') or 'Caption' in p.style.name:
                    pass  # keep scanning

    # Remove old table entries
    for el in to_remove:
        el.getparent().remove(el)
    print(f"  Removed {len(to_remove)} old LOT table entries")

    # Save after removal and reload
    doc2.save(DOC_PATH)
    doc3 = Document(DOC_PATH)
    paras3 = doc3.paragraphs

    # Find the last 'table of figures' entry (figures) after LOT heading
    # to insert after
    lot_h_idx = None
    last_fig_el = None
    for i, p in enumerate(paras3):
        t = p.text.strip()
        if ('List Of Tables' in t or 'LIST OF TABLES' in t) and lot_h_idx is None:
            lot_h_idx = i
        if lot_h_idx is not None and p.style.name == 'table of figures':
            if t.startswith('Figure') or not t.startswith('Table'):
                last_fig_el = p._element
            elif t.startswith('Table'):
                # Shouldn't be here - was cleaned above, but just in case
                last_fig_el = p._element

    # If no figure entries found, insert after the LOT heading itself
    if last_fig_el is None and lot_h_idx is not None:
        last_fig_el = paras3[lot_h_idx]._element
        print("  No figure entries found; inserting after LOT heading")

    def make_lot_entry(text, page_num):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), 'TableofFigures'); pPr.append(pStyle)
        p.append(pPr)
        p.append(make_run(text))
        # Tab
        r_tab = OxmlElement('w:r'); r_tab.append(OxmlElement('w:tab')); p.append(r_tab)
        p.append(make_run(str(page_num)))
        return p

    # Insert new LOT entries in order
    insert_after = last_fig_el
    for caption_text, _ in updated_captions:
        pg = caption_page_map.get(caption_text, '?')
        # LOT text = caption text without extra "Table X.Y: " duplication
        entry_el = make_lot_entry(caption_text, pg)
        insert_after.addnext(entry_el)
        insert_after = entry_el
        print(f"  Added LOT entry: {caption_text} -> p{pg}")

    doc3.save(DOC_PATH)
    print(f"\nDONE. Saved: {DOC_PATH}")
